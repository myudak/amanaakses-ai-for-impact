# -*- coding: utf-8 -*-
"""Generate the full academic AmanAkses proposal and AI documentation.

Format:
- A4
- Times New Roman, black
- 12 pt body, 1.5 line spacing
- margins: left 4 cm, right/top/bottom 3 cm
- editable image markers instead of embedded prototype screenshots
"""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
EDITABLE = ROOT / "supporting" / "editable"
OUTPUT = ROOT / "outputs" / "tugas-besar-ai"
ASSETS = ROOT / "supporting" / "assets"

EDITABLE.mkdir(parents=True, exist_ok=True)
OUTPUT.mkdir(parents=True, exist_ok=True)

LOGO = ASSETS / "undip-logo.png"
FONT = "Times New Roman"
BLACK = RGBColor(0, 0, 0)
LIGHT_GRAY = "E7E6E6"
VERY_LIGHT_GRAY = "F2F2F2"

MEMBERS = [
    ("Muchammad Yuda Tri Ananda", "24060124110142"),
    ("Nadia Azura Nurhaniya", "24060124120019"),
    ("Muhammad Zaidaan Ardiyansyah", "24060124140200"),
    ("Anintya Abhi Wiryateja", "24060124130053"),
    ("Muhamad Kemal Faza", "24060124120013"),
]

PROPOSAL_REFS = [
    "Komisi Nasional Anti Kekerasan terhadap Perempuan. (2024). Siaran Pers Hari Disabilitas Internasional 2024: Kuatkan Kepemimpinan Perempuan dengan Disabilitas untuk Masa Depan yang Lebih Inklusif dan Berkelanjutan. https://komnasperempuan.go.id/siaran-pers-detail/siaran-pers-komnas-perempuan-merespons-hari-disabilitas-internasional-2024",
    "Republik Indonesia. (2016). Undang-Undang Nomor 8 Tahun 2016 tentang Penyandang Disabilitas. https://peraturan.bpk.go.id/Home/Details/37251/uu-no-8-tahun-2016",
    "Republik Indonesia. (2022). Undang-Undang Nomor 12 Tahun 2022 tentang Tindak Pidana Kekerasan Seksual. https://peraturan.bpk.go.id/Details/207944/uu-no-12-tahun-2022",
    "Kementerian Pendidikan, Kebudayaan, Riset, dan Teknologi. (2024). Permendikbudristek Nomor 55 Tahun 2024 tentang Pencegahan dan Penanganan Kekerasan di Lingkungan Perguruan Tinggi. https://peraturan.bpk.go.id/Details/305767/permendikbudriset-no-55-tahun-",
    "Republik Indonesia. (2025). Peraturan Pemerintah Nomor 30 Tahun 2025 tentang Pencegahan Tindak Pidana Kekerasan Seksual serta Penanganan, Pelindungan, dan Pemulihan Korban. https://peraturan.bpk.go.id/Details/338353/pp-no-30-tahun-2025",
    "World Wide Web Consortium. (2023). Web Content Accessibility Guidelines (WCAG) 2.2. https://www.w3.org/TR/WCAG22/",
    "Google AI for Developers. (2026). Gemini API documentation: Structured output. https://ai.google.dev/gemini-api/docs/structured-output",
    "Tim Pengampu. (2026). Panduan Tugas Besar AI For Real Impact 2026: Solving Real Problems with AI Collaboration.",
]


def set_font(run, size=12, bold=False, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.color.rgb = BLACK
    run.bold = bold
    run.italic = italic
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), FONT)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=100, bottom=90, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def create_decimal_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(item.get(qn("w:abstractNumId")))
        for item in numbering.findall(qn("w:abstractNum"))
        if item.get(qn("w:abstractNumId")) is not None
    ]
    num_ids = [
        int(item.get(qn("w:numId")))
        for item in numbering.findall(qn("w:num"))
        if item.get(qn("w:numId")) is not None
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "720")
    indent.set(qn("w:hanging"), "360")
    p_pr.extend([tabs, indent])
    level.extend([start, num_fmt, lvl_text, lvl_jc, p_pr])
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])
    return run


def add_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(p, "PAGE")
    for run in p.runs:
        set_font(run, size=10)


def configure(doc):
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(4)
    section.right_margin = Cm(3)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(12)
    normal.font.color.rgb = BLACK
    normal.element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal.element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size, bold, before, after in (
        ("Heading 1", 12, True, 0, 12),
        ("Heading 2", 12, True, 10, 4),
        ("Heading 3", 12, True, 8, 3),
    ):
        style = doc.styles[style_name]
        style.font.name = FONT
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = BLACK
        style.element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style.element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = FONT
        style.font.size = Pt(12)
        style.font.color.rgb = BLACK
        style.element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style.element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        style.paragraph_format.space_after = Pt(3)

    add_page_number(section)


def paragraph(
    doc,
    text="",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    size=12,
    bold=False,
    italic=False,
    first_indent=1.0,
    space_before=0,
    space_after=6,
    keep_with_next=False,
):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.keep_with_next = keep_with_next
    if first_indent is not None:
        p.paragraph_format.first_line_indent = Cm(first_indent)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic)
    return p


def heading(doc, text, level=2):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_font(run, bold=True)
    return p


def chapter(doc, roman, title, page_break=True):
    p = doc.add_paragraph(style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.page_break_before = page_break
    p.paragraph_format.keep_with_next = True
    run = p.add_run(f"BAB {roman}")
    set_font(run, bold=True)
    run.add_break()
    run2 = p.add_run(title)
    set_font(run2, bold=True)
    return p


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.first_line_indent = Cm(-0.35)
        p.paragraph_format.keep_together = True
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(item)
        set_font(run)


def numbered(doc, items):
    for index, item in enumerate(items, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.keep_together = True
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.first_line_indent = Cm(-0.75)
        run = p.add_run(f"{index}. {item}")
        set_font(run)


def table(doc, headers, rows, widths_cm, font_size=10):
    table_obj = doc.add_table(rows=1, cols=len(headers))
    table_obj.style = "Table Grid"
    table_obj.autofit = False
    table_obj.alignment = 0

    header_row = table_obj.rows[0]
    set_repeat_table_header(header_row)
    prevent_row_split(header_row)
    for idx, header in enumerate(headers):
        cell = header_row.cells[idx]
        cell.width = Cm(widths_cm[idx])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_margins(cell)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.keep_with_next = True
        set_font(p.add_run(str(header)), size=font_size, bold=True)

    for row_values in rows:
        row = table_obj.add_row()
        prevent_row_split(row)
        for idx, value in enumerate(row_values):
            cell = row.cells[idx]
            cell.width = Cm(widths_cm[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            set_font(p.add_run(str(value)), size=font_size)
    paragraph(doc, "", first_indent=None, space_after=2)
    return table_obj


def caption(doc, text):
    return paragraph(
        doc,
        text,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=10,
        italic=True,
        first_indent=None,
        space_after=8,
        keep_with_next=False,
    )


def image_marker(doc, route_or_evidence, caption_text):
    p = paragraph(
        doc,
        f"[GAMBAR: {route_or_evidence}]",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=11,
        italic=True,
        first_indent=None,
        space_before=8,
        space_after=2,
        keep_with_next=True,
    )
    p.paragraph_format.keep_together = True
    caption(doc, caption_text)


def add_reference_list(doc, references):
    for ref in references:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.first_line_indent = Cm(-1)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(3)
        set_font(p.add_run(ref), size=10)


def cover(doc, document_type):
    paragraph(doc, document_type, WD_ALIGN_PARAGRAPH.CENTER, 14, True, first_indent=None, space_after=0)
    paragraph(doc, "TUGAS BESAR AI FOR REAL IMPACT 2026", WD_ALIGN_PARAGRAPH.CENTER, 14, True, first_indent=None, space_after=0)
    paragraph(
        doc,
        "AMANAKSES: PLATFORM DOKUMENTASI AMAN DAN AKSESIBEL\nBAGI PENYANDANG DISABILITAS",
        WD_ALIGN_PARAGRAPH.CENTER,
        14,
        True,
        first_indent=None,
        space_before=8,
        space_after=16,
    )
    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(14)
        p.add_run().add_picture(str(LOGO), width=Inches(1.8))
    paragraph(doc, "Disusun Oleh:", WD_ALIGN_PARAGRAPH.CENTER, 12, True, first_indent=None, space_after=4)
    for name, nim in MEMBERS:
        paragraph(doc, f"{name} ({nim})", WD_ALIGN_PARAGRAPH.CENTER, first_indent=None, space_after=0)
    paragraph(doc, "", first_indent=None, space_after=8)
    for line in (
        "DEPARTEMEN INFORMATIKA",
        "FAKULTAS SAINS DAN MATEMATIKA",
        "UNIVERSITAS DIPONEGORO",
        "2026",
    ):
        paragraph(doc, line, WD_ALIGN_PARAGRAPH.CENTER, 14, True, first_indent=None, space_after=0)
    doc.add_page_break()


def add_front_matter(doc, executive_summary, keywords, tables, figures):
    heading(doc, "RINGKASAN EKSEKUTIF", level=1)
    for text in executive_summary:
        paragraph(doc, text)
    paragraph(doc, f"Kata kunci: {keywords}", italic=True, first_indent=None)
    doc.add_page_break()

    heading(doc, "DAFTAR ISI", level=1)
    toc = doc.add_paragraph()
    toc.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    add_field(toc, 'TOC \\o "1-1" \\h \\z \\u')
    paragraph(
        doc,
        "Catatan: klik kanan daftar isi lalu pilih Update Field jika nomor halaman berubah setelah gambar diganti.",
        size=10,
        italic=True,
        first_indent=None,
        space_after=10,
    )
    p = paragraph(doc, "DAFTAR TABEL", size=10, bold=True, first_indent=None, space_before=2, space_after=1)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    for item in tables:
        p = paragraph(doc, item, size=10, first_indent=None, space_after=0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p = paragraph(doc, "DAFTAR GAMBAR", size=10, bold=True, first_indent=None, space_before=2, space_after=1)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    for item in figures:
        p = paragraph(doc, item, size=10, first_indent=None, space_after=0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def build_proposal():
    doc = Document()
    configure(doc)
    cover(doc, "PROPOSAL PROYEK")

    executive_summary = [
        "AmanAkses merupakan rancangan platform digital aksesibel yang membantu penyandang disabilitas memahami pilihan bantuan, mencatat pengalaman secara bertahap, menata bukti, menyusun kronologi, dan menyiapkan laporan awal tanpa kehilangan kendali atas data pribadi. Gagasan ini berangkat dari masalah nyata berupa hambatan akses informasi, komunikasi, dokumentasi, dan layanan ketika penyandang disabilitas mengalami atau berisiko mengalami kekerasan.",
        "Urgensi masalah diperkuat oleh laporan Komnas Perempuan yang mencatat 105 kasus kekerasan terhadap perempuan dengan disabilitas dalam CATAHU 2024, dengan 38 kasus dilaporkan langsung ke Komnas Perempuan. Angka tersebut tidak menggambarkan seluruh kejadian, tetapi menunjukkan bahwa kekerasan berbasis gender dan disabilitas tetap membutuhkan layanan yang inklusif. AmanAkses tidak diposisikan sebagai pengganti Satgas, pendamping, psikolog, penasihat hukum, atau layanan darurat. Platform dirancang sebagai ruang persiapan yang memberi waktu, pilihan, aksesibilitas, dan persetujuan yang jelas kepada pengguna.",
        "Pemanfaatan AI difokuskan pada Safe Timeline Assistant. Gemini hanya membantu mengubah catatan yang dipilih pengguna menjadi draf kronologi terstruktur. Setiap peristiwa harus memiliki referensi sumber, bagian yang ambigu atau hilang ditandai, dan pengguna wajib mengedit, menerima, atau menolak hasil. Prototype saat ini telah menunjukkan alur antarmuka utama, fitur timeline interaktif, fallback deterministik, serta pengujian otomatis. Fitur penyimpanan produksi, enkripsi, autentikasi, berbagi ke lembaga eksternal, dan integrasi layanan masih merupakan visi pengembangan.",
    ]
    proposal_tables = [
        "Tabel 2.1 Identitas kelompok",
        "Tabel 2.2 Persona dan kebutuhan akses utama",
        "Tabel 3.1 Pain point dalam perjalanan pengguna",
        "Tabel 3.2 Analisis akar masalah",
        "Tabel 4.1 Perbandingan pendekatan umum dan AmanAkses",
        "Tabel 4.2 Batas peran AI",
        "Tabel 5.1 Fitur, manfaat, dan status implementasi",
        "Tabel 5.2 Teknologi dan batas penggunaannya",
        "Tabel 7.1 Hasil validasi teknis",
        "Tabel 7.2 Risiko dan mitigasi",
        "Tabel 8.1 Roadmap pengembangan",
    ]
    proposal_figures = [
        "Gambar 5.1 Dashboard AmanAkses",
        "Gambar 5.2 Pahami Kekerasan",
        "Gambar 5.3 Jurnal Aman",
        "Gambar 5.4 Safe Timeline Assistant",
        "Gambar 5.5 Brankas Bukti",
        "Gambar 5.6 Pendamping Tepercaya",
        "Gambar 5.7 Laporan Awal",
        "Gambar 5.8 Pengaturan Aksesibilitas",
        "Gambar 5.9 Keluar Cepat",
    ]
    add_front_matter(
        doc,
        executive_summary,
        "aksesibilitas, penyandang disabilitas, dokumentasi aman, human-in-the-loop, Gemini, consent",
        proposal_tables,
        proposal_figures,
    )

    chapter(doc, "I", "PENDAHULUAN")
    heading(doc, "1.1 Latar Belakang", 2)
    paragraph(
        doc,
        "Kekerasan merupakan persoalan sosial yang tidak hanya berkaitan dengan terjadinya tindakan, tetapi juga dengan kemampuan korban untuk memahami situasi, memperoleh informasi, mendokumentasikan pengalaman, dan menjangkau bantuan. Bagi penyandang disabilitas, proses tersebut dapat menghadirkan hambatan tambahan. Informasi layanan sering menggunakan bahasa panjang dan teknis; formulir tidak selalu kompatibel dengan pembaca layar; pilihan komunikasi terbatas; tombol dan alur digital sulit dioperasikan; serta prosedur pelaporan dapat meminta pengguna menceritakan ulang pengalaman sensitif kepada beberapa pihak.",
    )
    paragraph(
        doc,
        "Komnas Perempuan dalam Siaran Pers Hari Disabilitas Internasional 2024 menyatakan bahwa perempuan dengan disabilitas masih menghadapi diskriminasi sistemik dan risiko kekerasan berbasis gender serta disabilitas. Berdasarkan CATAHU 2024, tercatat 105 kasus kekerasan terhadap perempuan dengan disabilitas dan 38 di antaranya dilaporkan langsung ke Komnas Perempuan. Data tersebut perlu dibaca secara hati-hati karena kasus yang tercatat tidak selalu mewakili seluruh kejadian. Meskipun demikian, data tersebut cukup menunjukkan adanya kebutuhan terhadap mekanisme informasi, pendokumentasian, dan rujukan yang lebih inklusif.",
    )
    paragraph(
        doc,
        "Secara normatif, Undang-Undang Nomor 8 Tahun 2016 menjamin pemenuhan hak penyandang disabilitas, termasuk aksesibilitas dan akomodasi yang layak. Undang-Undang Nomor 12 Tahun 2022 menegaskan bahwa korban penyandang disabilitas berhak memperoleh aksesibilitas dan akomodasi yang layak dalam pemenuhan haknya. PP Nomor 30 Tahun 2025 juga menyatakan bahwa pencegahan tindak pidana kekerasan seksual harus memenuhi aksesibilitas bagi anak, penyandang disabilitas, dan lanjut usia. Dalam konteks perguruan tinggi, Permendikbudristek Nomor 55 Tahun 2024 mengatur pencegahan, penanganan, pemulihan, hak korban dan saksi, satuan tugas, serta pengelolaan data kekerasan.",
    )
    paragraph(
        doc,
        "Ketersediaan regulasi belum otomatis menghilangkan hambatan operasional pada tingkat pengguna. Seseorang masih perlu memahami pilihan, mengingat detail, menyusun urutan waktu, mengelola bukti, menentukan pihak yang dipercaya, dan memutuskan informasi apa yang ingin dibagikan. Ketika proses digital tidak aksesibel atau terlalu memaksa, pengguna dapat kehilangan energi, rasa aman, dan kendali. Pendekatan yang hanya menyediakan formulir pelaporan belum tentu menjawab kebutuhan untuk belajar, berhenti sejenak, menyimpan draf, atau meminta pendampingan sebelum mengambil keputusan.",
    )
    paragraph(
        doc,
        "Berdasarkan kesenjangan tersebut, tim mengembangkan AmanAkses sebagai visi platform dokumentasi aman dan aksesibel. AmanAkses menghubungkan edukasi easy-read, jurnal bertahap, brankas bukti, penyusunan kronologi, persetujuan berbagi, pratinjau laporan, direktori bantuan, pengaturan aksesibilitas, dan fitur keluar cepat. AI ditempatkan pada satu fungsi terbatas, yaitu membantu menata catatan menjadi draf kronologi. Dengan demikian, proyek tidak dimulai dari keinginan menggunakan AI, melainkan dari masalah nyata yang kemudian dipertemukan dengan fungsi AI yang relevan dan dapat dibatasi.",
    )

    heading(doc, "1.2 Identifikasi Masalah", 2)
    bullets(
        doc,
        [
            "Informasi mengenai kekerasan, hak, dan pilihan bantuan belum selalu tersedia dalam bentuk yang mudah dibaca, mudah dinavigasi, dan setara bagi berbagai ragam disabilitas.",
            "Catatan pengalaman dan bukti sering tersimpan terpisah sehingga sulit disusun kembali menjadi kronologi yang konsisten.",
            "Proses pelaporan dapat menuntut pengguna mengulang cerita, meningkatkan beban emosional dan risiko kehilangan detail.",
            "Pengguna belum selalu memperoleh penjelasan granular tentang data apa yang dibagikan, kepada siapa, untuk berapa lama, dan bagaimana mencabut akses.",
            "AI generatif dapat membantu merapikan informasi, tetapi berisiko menambah fakta, menghapus nuansa, atau menghasilkan kesimpulan yang tidak didukung sumber.",
        ],
    )
    heading(doc, "1.3 Rumusan Masalah", 2)
    numbered(
        doc,
        [
            "Bagaimana merancang alur belajar, mencatat, menata bukti, dan menyiapkan laporan yang aksesibel serta tidak memaksa?",
            "Bagaimana mengurangi beban pengguna dalam menyusun kronologi tanpa menjadikan AI sebagai penentu kebenaran atau pembuat keputusan?",
            "Bagaimana mempertahankan kendali pengguna atas pemilihan catatan, peninjauan hasil, persetujuan berbagi, dan penghentian proses?",
            "Bagaimana membedakan secara jujur antara visi produk, kemampuan prototype, dan fitur yang masih harus dikembangkan?",
        ],
    )
    heading(doc, "1.4 Tujuan", 2)
    numbered(
        doc,
        [
            "Menghasilkan prototype web yang menunjukkan alur dukungan aksesibel dari edukasi hingga pratinjau laporan awal.",
            "Menerapkan Safe Timeline Assistant untuk menyusun draf peristiwa yang bersumber, dapat diedit, dan wajib ditinjau manusia.",
            "Menerapkan prinsip accessibility-by-design, privacy-by-design, consent-by-design, dan human-in-the-loop pada rancangan solusi.",
            "Memvalidasi fungsi inti menggunakan data sintetis, pengujian struktur respons, fallback, build, lint, dan pemeriksaan aksesibilitas dasar.",
        ],
    )
    heading(doc, "1.5 Manfaat", 2)
    bullets(
        doc,
        [
            "Bagi pengguna: menyediakan ruang untuk memahami situasi dan menyusun informasi secara bertahap sesuai kesiapan.",
            "Bagi pendamping: menyediakan bahan awal yang lebih terstruktur tanpa menghapus konteks atau kebutuhan akses pengguna.",
            "Bagi lembaga layanan: memberikan rancangan laporan awal yang mencantumkan kronologi, bukti terpilih, kebutuhan komunikasi, dan persetujuan.",
            "Bagi tim akademik: menjadi studi kasus pemanfaatan AI yang terukur, kritis, dan tidak berorientasi pada kecanggihan semata.",
        ],
    )
    heading(doc, "1.6 Urgensi dan Relevansi dengan AI For Real Impact", 2)
    paragraph(
        doc,
        "AmanAkses relevan dengan tema AI For Real Impact karena berangkat dari kelompok pengguna yang jelas, masalah yang memiliki dampak, serta kebutuhan akan prototype yang dapat didemokan. Peran AI tidak diperluas secara artifisial. Gemini digunakan untuk tugas transformasi informasi yang spesifik, sedangkan keputusan tetap berada pada pengguna. Strategi ini mendukung tujuan pembelajaran berupa problem solving, kolaborasi kritis dengan AI, validasi hasil, dan pengembangan solusi yang berpotensi dilanjutkan ke kompetisi, penelitian, atau pengabdian masyarakat.",
    )

    chapter(doc, "II", "IDENTITAS KELOMPOK DAN TARGET PENGGUNA", page_break=False)
    heading(doc, "2.1 Identitas Kelompok", 2)
    table(
        doc,
        ["No.", "Nama", "NIM"],
        [(index, name, nim) for index, (name, nim) in enumerate(MEMBERS, start=1)],
        [1.2, 8.3, 4.7],
        11,
    )
    caption(doc, "Tabel 2.1 Identitas kelompok")
    heading(doc, "2.2 Target Pengguna Utama dan Sekunder", 2)
    paragraph(
        doc,
        "Target utama AmanAkses adalah mahasiswa atau masyarakat penyandang disabilitas yang memerlukan ruang dokumentasi bertahap dan aksesibel. Target sekunder meliputi pendamping tepercaya, Satgas Pencegahan dan Penanganan Kekerasan di perguruan tinggi, lembaga layanan, bantuan hukum, serta pihak lain yang hanya menerima informasi berdasarkan persetujuan pengguna.",
    )
    heading(doc, "2.3 Persona dan Kebutuhan Akses", 2)
    table(
        doc,
        ["Persona", "Hambatan yang Mungkin Dialami", "Kebutuhan Desain"],
        [
            ("Netra atau low vision", "Struktur halaman tidak semantik, kontras rendah, kontrol tanpa label.", "Screen reader, keyboard, heading jelas, teks alternatif, kontras dan ukuran teks."),
            ("Tuli atau hambatan pendengaran", "Informasi hanya berbentuk audio dan tidak ada padanan teks.", "Konten tertulis, caption, indikator visual, bahasa isyarat pada pengembangan lanjutan."),
            ("Disabilitas fisik", "Target klik kecil dan interaksi memerlukan gerak presisi.", "Kontrol besar, keyboard, langkah singkat, pilihan input suara pada visi produk."),
            ("Disabilitas intelektual atau kognitif", "Bahasa kompleks, terlalu banyak pilihan, dan beban mengingat urutan.", "Easy-read, satu keputusan per langkah, pengulangan konfirmasi, draf bertahap."),
            ("Disabilitas psikososial", "Alur memaksa, tidak ada jeda, dan tampilan memicu rasa tidak aman.", "Bahasa suportif, skip, autosave, safe exit, dan kendali penuh atas kelanjutan proses."),
            ("Pendamping atau Satgas", "Informasi awal tidak terstruktur dan kebutuhan akses tidak tercatat.", "Ringkasan yang dapat ditinjau, sumber yang jelas, kebutuhan komunikasi, dan status consent."),
        ],
        [3.4, 5.0, 5.8],
        9,
    )
    caption(doc, "Tabel 2.2 Persona dan kebutuhan akses utama")
    heading(doc, "2.4 Prinsip Hubungan dengan Pengguna", 2)
    bullets(
        doc,
        [
            "Tidak memaksa pengguna menyelesaikan seluruh alur dalam satu sesi.",
            "Tidak menggunakan bahasa yang menyalahkan korban atau mengklaim kebenaran kasus.",
            "Memberi pilihan untuk melewati fitur AI dan mengedit hasil secara manual.",
            "Menjelaskan status simulasi, keterbatasan prototype, penerima data, dan konsekuensi setiap tindakan.",
            "Menggunakan data sintetis pada pengembangan, pengujian, screenshot, dan demonstrasi kelas.",
        ],
    )

    chapter(doc, "III", "ANALISIS MASALAH DAN KEBUTUHAN", page_break=False)
    heading(doc, "3.1 Perjalanan Pengguna Saat Ini", 2)
    paragraph(
        doc,
        "Perjalanan pengguna tidak selalu dimulai dari keputusan untuk melapor. Pengguna dapat berada pada tahap memahami apakah suatu situasi membuatnya tidak aman, mencari informasi, mencatat sedikit demi sedikit, meminta pendapat pendamping, menata bukti, atau baru mempertimbangkan layanan formal. Karena itu, solusi yang langsung mengarahkan pengguna ke formulir final berisiko tidak sesuai dengan kesiapan dan kebutuhan aksesnya.",
    )
    table(
        doc,
        ["Tahap", "Aktivitas Pengguna", "Pain Point", "Peluang Perbaikan"],
        [
            ("Memahami", "Mencari informasi mengenai bentuk kekerasan dan hak.", "Bahasa rumit, materi tidak aksesibel, takut salah memahami.", "Materi easy-read dan navigasi aksesibel."),
            ("Mencatat", "Menyimpan detail yang masih diingat.", "Catatan tercecer, takut terlihat orang lain, sulit menulis panjang.", "Jurnal bertahap, autosave, dan mode privat."),
            ("Menata", "Menghubungkan catatan, waktu, lokasi, dan bukti.", "Beban kognitif tinggi dan informasi hilang.", "Draf kronologi bersumber yang tidak menebak data."),
            ("Meminta bantuan", "Memilih pihak tepercaya dan menjelaskan kebutuhan.", "Harus bercerita ulang dan kebutuhan akses terlewat.", "Ringkasan terpilih dan profil aksesibilitas."),
            ("Berbagi", "Mengirim laporan atau bukti.", "Tidak jelas siapa melihat apa dan berapa lama.", "Consent granular dan pratinjau sebelum berbagi."),
        ],
        [2.3, 4.0, 4.2, 3.7],
        8,
    )
    caption(doc, "Tabel 3.1 Pain point dalam perjalanan pengguna")
    heading(doc, "3.2 Analisis Akar Masalah", 2)
    table(
        doc,
        ["Dimensi", "Akar Masalah", "Dampak"],
        [
            ("Aksesibilitas", "Layanan digital dirancang untuk pengguna rata-rata dan aksesibilitas ditambahkan belakangan.", "Sebagian pengguna tidak dapat menyelesaikan alur secara mandiri."),
            ("Informasi", "Konten tersebar, teknis, dan tidak disajikan berdasarkan tahap kesiapan.", "Pengguna kesulitan memahami pilihan dan konsekuensi."),
            ("Dokumentasi", "Catatan dan bukti tidak memiliki struktur atau konteks bersama.", "Detail mudah hilang dan kronologi sulit disusun."),
            ("Kepercayaan", "Proses berbagi data kurang transparan dan sulit dibatalkan.", "Pengguna ragu menggunakan layanan atau membagikan terlalu banyak data."),
            ("AI", "Model generatif cenderung menghasilkan jawaban yang meyakinkan meskipun sumber tidak lengkap.", "Risiko fakta tambahan, bias, dan keputusan yang tidak dapat diaudit."),
        ],
        [2.7, 7.0, 4.5],
        9,
    )
    caption(doc, "Tabel 3.2 Analisis akar masalah")
    heading(doc, "3.3 Kesenjangan Solusi yang Tersedia", 2)
    paragraph(
        doc,
        "Catatan ponsel membantu menyimpan teks, tetapi tidak menggabungkan kebutuhan akses, bukti, consent, dan laporan. Folder screenshot membantu menyimpan media, tetapi konteks, sumber, serta keterkaitan antarperistiwa mudah hilang. Form aduan menyediakan kanal formal, tetapi sering mengasumsikan pengguna sudah siap melapor dan mampu mengisi seluruh informasi. Chatbot umum mudah diakses, tetapi dapat memberikan jawaban terlalu luas, menambah interpretasi, atau tidak menyediakan mekanisme audit sumber.",
    )
    heading(doc, "3.4 Kebutuhan Fungsional Prioritas", 2)
    numbered(
        doc,
        [
            "Menyajikan informasi dasar dan opsi bantuan dalam bahasa sederhana.",
            "Memungkinkan pengguna membuat dan menyimpan catatan secara bertahap.",
            "Memungkinkan pemilihan catatan secara eksplisit sebelum diproses AI.",
            "Menghasilkan draf kronologi dengan sumber dan status ketidakpastian.",
            "Mewajibkan pengguna mengedit, menerima, atau menolak setiap peristiwa.",
            "Menampilkan pratinjau laporan dan ruang lingkup data sebelum berbagi.",
            "Menyediakan pengaturan aksesibilitas dan keluar cepat pada seluruh alur sensitif.",
        ],
    )
    heading(doc, "3.5 Kebutuhan Non-Fungsional", 2)
    bullets(
        doc,
        [
            "Aksesibilitas: mengikuti prinsip perceivable, operable, understandable, dan robust pada WCAG 2.2.",
            "Privasi: meminimalkan data, menghindari data nyata pada demo, dan tidak menaruh rahasia pada browser.",
            "Keandalan: menyediakan fallback agar fungsi demonstrasi tetap berjalan ketika layanan model gagal.",
            "Auditabilitas: setiap peristiwa memiliki sumber dan status review.",
            "Keamanan emosional: bahasa netral, tidak grafis, tidak memaksa, serta mendukung jeda dan penghentian proses.",
        ],
    )

    chapter(doc, "IV", "IDE SOLUSI DAN NILAI INOVASI", page_break=False)
    heading(doc, "4.1 Konsep AmanAkses", 2)
    paragraph(
        doc,
        "AmanAkses dirancang sebagai ruang persiapan digital, bukan sekadar kanal aduan. Nilai utamanya dirumuskan sebagai Understand safely, Record safely, Decide safely, dan Share safely. Pengguna dapat memulai dari materi edukasi, mencatat saat siap, memilih informasi yang relevan, meminta bantuan AI secara opsional, meninjau hasil, menentukan pendamping, dan memutuskan apakah laporan akan dilanjutkan.",
    )
    heading(doc, "4.2 Keunikan Solusi", 2)
    bullets(
        doc,
        [
            "Menggabungkan aksesibilitas multi-disabilitas dengan alur dokumentasi dan consent.",
            "Menempatkan pengguna sebagai pemilik keputusan, bukan objek yang harus mengikuti alur tunggal.",
            "Menggunakan AI untuk pekerjaan administratif-kognitif yang sempit, bukan untuk menilai kasus.",
            "Menampilkan catatan sumber dan ketidakpastian sehingga draf dapat diaudit.",
            "Menyediakan jalur fallback dan jalur manual sehingga manfaat produk tidak sepenuhnya bergantung pada model.",
        ],
    )
    table(
        doc,
        ["Pendekatan Umum", "Keterbatasan", "Pendekatan AmanAkses"],
        [
            ("Catatan biasa", "Tidak terstruktur dan sulit disatukan dengan bukti.", "Jurnal, sumber, timeline, dan laporan berada dalam satu alur."),
            ("Form laporan langsung", "Mengasumsikan pengguna siap serta mampu mengisi sekaligus.", "Pengguna dapat belajar, mencatat, berhenti, dan melanjutkan."),
            ("Berbagi folder", "Ruang lingkup data dan durasi akses tidak jelas.", "Visi consent granular dan pratinjau sebelum berbagi."),
            ("Chatbot umum", "Dapat memberi kesimpulan luas tanpa sumber.", "AI dibatasi pada structured output, sumber, dan review per event."),
        ],
        [3.3, 5.2, 5.7],
        9,
    )
    caption(doc, "Tabel 4.1 Perbandingan pendekatan umum dan AmanAkses")
    heading(doc, "4.3 Peran Safe Timeline Assistant", 2)
    paragraph(
        doc,
        "Safe Timeline Assistant menerima hanya catatan yang dipilih pengguna. Fungsi serverless membentuk permintaan terstruktur kepada Gemini, memeriksa respons, lalu mengembalikan draf peristiwa. Setiap peristiwa memuat tanggal, waktu, lokasi, ringkasan, ID sumber, jenis ketidakpastian, dan penanda wajib review. Jika Gemini tidak tersedia, fallback deterministik menyusun hasil secara konservatif dari label eksplisit pada catatan sintetis.",
    )
    table(
        doc,
        ["AI Boleh", "AI Tidak Boleh"],
        [
            ("Mengambil tanggal, waktu, dan lokasi yang tertulis.", "Mengisi fakta yang hilang dari dugaan atau metadata lain."),
            ("Meringkas catatan dengan bahasa netral.", "Menilai apakah laporan benar atau salah."),
            ("Menyertakan ID catatan sumber.", "Menyimpulkan niat, kesalahan, identitas, atau motif."),
            ("Menandai explicit, ambiguous, atau missing.", "Memberikan diagnosis atau keputusan hukum."),
            ("Menghasilkan draf untuk diperiksa.", "Mengirim laporan atau bukti secara otomatis."),
        ],
        [7.1, 7.1],
        10,
    )
    caption(doc, "Tabel 4.2 Batas peran AI")
    heading(doc, "4.4 Alur End-to-End", 2)
    numbered(
        doc,
        [
            "Pengguna mengatur kebutuhan aksesibilitas dan membuka ruang pribadi.",
            "Pengguna mempelajari informasi dasar tanpa kewajiban membuat laporan.",
            "Pengguna menulis jurnal dan memilih catatan yang relevan.",
            "Safe Timeline Assistant menghasilkan draf kronologi bersumber.",
            "Pengguna membuka sumber, mengedit, menerima, atau menolak setiap peristiwa.",
            "Pengguna memilih bukti dan kebutuhan akses yang akan disertakan.",
            "Pengguna meninjau laporan awal dan memberikan persetujuan eksplisit sebelum langkah eksternal.",
        ],
    )

    chapter(doc, "V", "FITUR, TEKNOLOGI, DAN BATAS IMPLEMENTASI")
    heading(doc, "5.1 Fitur dan Status Prototype", 2)
    table(
        doc,
        ["Fitur", "Nilai bagi Pengguna", "Status Saat Ini"],
        [
            ("Dashboard", "Akses ringkas ke seluruh alur dan status demo.", "Interaktif"),
            ("Pahami Kekerasan", "Materi easy-read dan opsi belajar bertahap.", "Interaktif berbasis konten sintetis"),
            ("Jurnal Aman", "Pencatatan bertahap dan pemilihan catatan.", "UI interaktif; penyimpanan produksi belum ada"),
            ("Safe Timeline Assistant", "Draf kronologi bersumber dan dapat diedit.", "Interaktif; Gemini serverless dan fallback"),
            ("Brankas Bukti", "Metadata, tag, hash, dan konteks bukti.", "Simulasi UI"),
            ("Pendamping Tepercaya", "Visi pemilihan scope akses dan consent.", "Simulasi UI"),
            ("Laporan Awal", "Ringkasan untuk ditinjau manusia.", "Preview simulasi"),
            ("Aksesibilitas", "Ukuran teks, kontras, kontrol besar, reduced motion.", "Interaktif pada state demo"),
            ("Keluar Cepat", "Mengalihkan pengguna dari halaman sensitif.", "Interaktif"),
        ],
        [3.4, 6.3, 4.3],
        8,
    )
    caption(doc, "Tabel 5.1 Fitur, manfaat, dan status implementasi")

    image_marker(doc, "Gambar web /app/dashboard", "Gambar 5.1 Dashboard AmanAkses")
    image_marker(doc, "Gambar web /app/pahami-kekerasan", "Gambar 5.2 Pahami Kekerasan")
    image_marker(doc, "Gambar web /app/jurnal", "Gambar 5.3 Jurnal Aman")
    image_marker(doc, "Gambar web /app/kronologi", "Gambar 5.4 Safe Timeline Assistant")
    image_marker(doc, "Gambar web /app/brankas-bukti", "Gambar 5.5 Brankas Bukti")
    image_marker(doc, "Gambar web /app/pendamping", "Gambar 5.6 Pendamping Tepercaya")
    image_marker(doc, "Gambar web /app/laporan", "Gambar 5.7 Laporan Awal")
    image_marker(doc, "Gambar web /app/aksesibilitas", "Gambar 5.8 Pengaturan Aksesibilitas")
    image_marker(doc, "Gambar web /safe-exit", "Gambar 5.9 Keluar Cepat")

    heading(doc, "5.2 Arsitektur Teknis Prototype", 2)
    paragraph(
        doc,
        "Prototype menggunakan React, TypeScript, Vite, React Router, Tailwind CSS, dan Motion pada sisi antarmuka. Permintaan timeline dikirim ke POST /api/timeline yang berjalan sebagai fungsi serverless Vercel. GEMINI_API_KEY dibaca hanya pada sisi server. Respons Gemini dibatasi menggunakan structured JSON dan divalidasi kembali oleh aplikasi. UI kemudian menampilkan hasil sebagai draf yang dapat diedit dan diputuskan oleh pengguna.",
    )
    table(
        doc,
        ["Lapisan", "Teknologi", "Fungsi dan Batas"],
        [
            ("Frontend", "React 19, TypeScript, Vite", "Menampilkan seluruh demo; tidak boleh menerima API key."),
            ("Navigasi/UI", "React Router, Tailwind CSS, Motion", "Alur SPA, aksesibilitas dasar, dan animasi yang dapat dikurangi."),
            ("Endpoint", "Vercel serverless function", "Memvalidasi input, menghubungi Gemini, dan menjaga rahasia di server."),
            ("AI runtime", "Gemini structured output", "Menyusun draf timeline; bukan sumber kebenaran."),
            ("Fallback", "Ekstraksi deterministik", "Menjaga demo berjalan tanpa jaringan/API; bukan NLP lengkap."),
            ("Validasi", "TypeScript schema checks dan fixtures", "Menolak struktur atau sumber yang tidak valid."),
        ],
        [2.7, 4.4, 6.9],
        9,
    )
    caption(doc, "Tabel 5.2 Teknologi dan batas penggunaannya")
    heading(doc, "5.3 Visi Arsitektur Produksi", 2)
    paragraph(
        doc,
        "Pada pengembangan produksi, AmanAkses memerlukan autentikasi yang aman, penyimpanan terenkripsi, object storage untuk bukti, consent ledger, audit log, kebijakan retensi, penghapusan data, pengelolaan layanan bantuan, dan pengujian keamanan independen. Komponen tersebut merupakan rencana produk dan tidak diklaim telah tersedia pada prototype tugas besar.",
    )
    heading(doc, "5.4 Aksesibilitas, Privasi, dan Consent", 2)
    bullets(
        doc,
        [
            "Aksesibilitas diperlakukan sebagai kebutuhan arsitektural, bukan lapisan kosmetik.",
            "Data yang dikirim ke AI harus minimum dan dipilih secara eksplisit.",
            "Catatan tidak boleh dicatat ke log aplikasi atau digunakan untuk pelatihan model tanpa dasar dan persetujuan yang jelas.",
            "Setiap hasil AI diberi label draf, menampilkan sumber, dan menyediakan kontrol edit/terima/tolak.",
            "Tidak ada auto-submit; pembagian ke pihak eksternal merupakan visi yang harus selalu diawali pratinjau dan consent.",
        ],
    )

    chapter(doc, "VI", "RENCANA PROTOTYPE DAN SKENARIO DEMONSTRASI", page_break=False)
    heading(doc, "6.1 Bentuk Prototype", 2)
    paragraph(
        doc,
        "Prototype berbentuk aplikasi web responsif yang dapat dijalankan secara lokal dan disiapkan untuk deployment Vercel. Seluruh data yang tersedia di dalam aplikasi merupakan data sintetis. Prototype diprioritaskan untuk menunjukkan hubungan antara masalah, alur pengguna, fungsi AI, tinjauan manusia, dan batasan produk; bukan untuk memproses data penyintas nyata.",
    )
    heading(doc, "6.2 Skenario Demo Utama", 2)
    numbered(
        doc,
        [
            "Buka dashboard dan jelaskan bahwa pengguna dapat memulai dari belajar, mencatat, atau mencari bantuan.",
            "Tunjukkan Pahami Kekerasan dan aksesibilitas sebagai konteks sebelum pelaporan.",
            "Buka Jurnal Aman dan pilih catatan sintetis yang boleh diproses.",
            "Buka Safe Timeline Assistant, susun draf, dan jelaskan mode Gemini atau fallback.",
            "Periksa referensi sumber, edit satu peristiwa, terima satu peristiwa, dan tolak satu peristiwa.",
            "Buka Brankas Bukti, Pendamping, dan Laporan Awal sebagai gambaran alur lanjutan.",
            "Tunjukkan pengaturan aksesibilitas dan fitur Keluar Cepat.",
        ],
    )
    heading(doc, "6.3 Rencana Penyempurnaan Setelah Tugas Besar", 2)
    bullets(
        doc,
        [
            "Melakukan co-design bersama organisasi disabilitas, pendamping, dan pengelola layanan.",
            "Menguji dukungan screen reader, keyboard, zoom, kontras, caption, easy-read, dan bahasa isyarat.",
            "Membangun autentikasi, penyimpanan terenkripsi, consent ledger, dan revoke access.",
            "Mengevaluasi model AI dengan dataset sintetis yang lebih beragam dan red-team prompt.",
            "Melakukan pilot terbatas setelah penilaian etika, keamanan, privasi, dan kesiapan layanan.",
        ],
    )

    chapter(doc, "VII", "VALIDASI, EVALUASI, DAN MANAJEMEN RISIKO", page_break=False)
    heading(doc, "7.1 Validasi yang Telah Dilakukan", 2)
    table(
        doc,
        ["Pemeriksaan", "Tujuan", "Hasil"],
        [
            ("Build TypeScript dan Vite", "Memastikan prototype dapat dikompilasi sebagai build produksi.", "Lulus"),
            ("ESLint", "Mendeteksi masalah statis pada source code.", "Lulus"),
            ("10 fixture utama", "Memeriksa field eksplisit, field hilang, ambiguitas, alias, dan catatan netral.", "10/10 lulus"),
            ("1 negative guard", "Menolak respons dengan ID sumber yang tidak terdapat pada input.", "Lulus"),
            ("Fallback failure path", "Memastikan demo tetap menghasilkan draf ketika endpoint/API gagal.", "Lulus"),
            ("Secret audit", "Memastikan API key tidak masuk bundle browser atau arsip source.", "Lulus"),
            ("Keyboard dan aksesibilitas dasar", "Memeriksa label pencarian, bahasa dokumen, navigasi, dan alur utama.", "Diperiksa"),
            ("Peer testing", "Menguji pemahaman dan kemudahan alur oleh 3-5 mahasiswa.", "[[ISI_HASIL_UJI_PEER]]"),
        ],
        [4.0, 7.0, 3.2],
        9,
    )
    caption(doc, "Tabel 7.1 Hasil validasi teknis")
    paragraph(
        doc,
        "Secara ringkas, hasil validasi tersebut dinyatakan sebagai 10 skenario utama lulus + 1 negative guard lulus. Rumusan ini membedakan fixture perilaku utama dari pengujian penolakan respons yang tidak memiliki referensi sumber valid.",
    )
    heading(doc, "7.2 Interpretasi Hasil", 2)
    paragraph(
        doc,
        "Hasil validasi menunjukkan bahwa kontrak data, fallback, dan guard sumber bekerja pada skenario sintetis yang dirancang. Hasil tersebut belum membuktikan bahwa model selalu benar pada data dunia nyata. Structured output menjamin bentuk respons, bukan ketepatan makna. Oleh karena itu, sumber, ketidakpastian, dan keputusan pengguna tetap menjadi bagian wajib dari antarmuka.",
    )
    heading(doc, "7.3 Rencana Peer Testing", 2)
    paragraph(
        doc,
        "Peer testing direncanakan kepada 3-5 mahasiswa menggunakan data sintetis. Peserta diminta menyelesaikan alur memilih catatan, membuat draf timeline, membuka sumber, mengedit, menerima atau menolak, dan memahami status laporan. Observasi berfokus pada keterbacaan, kejelasan label draf AI, pemahaman consent, navigasi keyboard, serta bagian yang membingungkan.",
    )
    bullets(
        doc,
        [
            "Jumlah peserta: [[ISI_JUMLAH_PESERTA_PEER]].",
            "Tanggal pengujian: [[ISI_TANGGAL_UJI_PEER]].",
            "Temuan utama: [[ISI_TEMUAN_UJI_PEER]].",
            "Revisi setelah pengujian: [[ISI_REVISI_SETELAH_UJI_PEER]].",
        ],
    )
    heading(doc, "7.4 Risiko dan Mitigasi", 2)
    table(
        doc,
        ["Risiko", "Dampak", "Mitigasi"],
        [
            ("AI menambah fakta", "Kronologi menyesatkan.", "Prompt larangan, field nullable, sumber wajib, validator, dan review manusia."),
            ("Bahasa AI menghapus nuansa", "Pengalaman pengguna tereduksi.", "Tampilkan catatan sumber dan izinkan edit manual penuh."),
            ("Data sensitif bocor", "Kerugian privasi dan keselamatan.", "Data sintetis pada demo, key server-side, minimisasi data, dan no logging."),
            ("Pengguna mengira produk layanan darurat", "Terlambat mencari bantuan langsung.", "Disclaimer dan pusat bantuan; produk bukan layanan respons real-time."),
            ("Aksesibilitas tidak sesuai kebutuhan nyata", "Pengguna tetap mengalami hambatan.", "Co-design dan pengujian dengan ragam disabilitas sebelum produksi."),
            ("Visi disalahartikan sebagai fitur selesai", "Ekspektasi dan evaluasi tidak akurat.", "Status interaktif, simulasi, dan rencana ditulis pada dokumen dan UI."),
        ],
        [4.0, 4.5, 5.7],
        9,
    )
    caption(doc, "Tabel 7.2 Risiko dan mitigasi")
    heading(doc, "7.5 Keterbatasan", 2)
    bullets(
        doc,
        [
            "Data pengembangan dan pengujian seluruhnya sintetis.",
            "Prototype belum diuji langsung bersama penyandang disabilitas atau penyintas.",
            "Penyimpanan, enkripsi, autentikasi, upload bukti, dan berbagi eksternal masih simulasi atau visi.",
            "Gemini live bergantung pada konfigurasi API dan dapat menghasilkan kesalahan semantik.",
            "Direktori bantuan dan identitas layanan pada demo bukan data operasional.",
        ],
    )

    chapter(doc, "VIII", "DAMPAK, ROADMAP, DAN PENUTUP", page_break=False)
    heading(doc, "8.1 Dampak yang Diharapkan", 2)
    paragraph(
        doc,
        "Dampak jangka pendek AmanAkses adalah membantu pengguna mengurangi beban mengingat dan menata informasi. Dampak menengah adalah menyediakan cara komunikasi yang lebih terstruktur kepada pendamping atau layanan. Dampak jangka panjang hanya dapat dicapai jika produk dikembangkan bersama komunitas, memiliki tata kelola data yang kuat, dan terhubung dengan layanan yang benar-benar siap menerima pengguna secara aksesibel.",
    )
    heading(doc, "8.2 Roadmap Pengembangan", 2)
    table(
        doc,
        ["Tahap", "Fokus", "Keluaran"],
        [
            ("1. Prototype kelas", "UI utama, timeline AI, fallback, fixture, dokumentasi.", "Demo dan paket tugas besar."),
            ("2. Validasi partisipatif", "Co-design, peer testing, audit aksesibilitas, penyempurnaan bahasa.", "Temuan kebutuhan dan backlog terprioritas."),
            ("3. Fondasi keamanan", "Autentikasi, enkripsi, consent ledger, audit log, retensi.", "MVP terbatas dengan data terlindungi."),
            ("4. Pilot layanan", "Integrasi lembaga, SOP, pelatihan, monitoring, dan incident response.", "Pilot terbatas yang dievaluasi secara etis."),
            ("5. Pengembangan lanjutan", "Penelitian, GEMASTIK, PKM, pengabdian, dan evaluasi dampak.", "Produk dan bukti dampak yang lebih matang."),
        ],
        [2.9, 6.3, 5.0],
        9,
    )
    caption(doc, "Tabel 8.1 Roadmap pengembangan")
    heading(doc, "8.3 Penutup", 2)
    paragraph(
        doc,
        "AmanAkses menunjukkan bahwa pemanfaatan AI yang bermakna tidak harus dimulai dari model yang kompleks. Nilai utama proyek terletak pada pemahaman masalah, penetapan batas, desain aksesibel, dan mekanisme agar pengguna tetap memegang keputusan. Safe Timeline Assistant membantu menata, tetapi tidak menentukan. Dengan validasi yang jujur dan roadmap yang bertahap, AmanAkses memiliki potensi untuk dikembangkan menjadi solusi yang lebih matang melalui kolaborasi dengan komunitas disabilitas, lembaga layanan, peneliti, dan pengembang.",
    )
    heading(doc, "DAFTAR PUSTAKA", 1)
    add_reference_list(doc, PROPOSAL_REFS)

    path = EDITABLE / "Proposal_AmanAkses_Akademik.docx"
    doc.save(path)
    return path


def build_ai_doc():
    doc = Document()
    configure(doc)
    cover(doc, "DOKUMENTASI PENGGUNAAN AI")

    ai_summary = [
        "Dokumen ini menjelaskan penggunaan AI secara lengkap selama pengembangan AmanAkses. Dokumentasi dibagi menjadi dua kategori yang tidak boleh dicampur: AI di dalam produk dan AI sebagai partner kerja tim. AI di dalam produk adalah Gemini pada Safe Timeline Assistant. AI sebagai partner kerja tim adalah ChatGPT dan Codex untuk brainstorming, analisis, desain solusi, implementasi, debugging, pengujian, dan penyusunan dokumen.",
        "Python-docx, Microsoft Word, PowerPoint, Vite, ESLint, dan perangkat render bukan AI. Alat-alat tersebut digunakan untuk produksi artefak, kompilasi, dan pemeriksaan. Pemisahan ini penting agar dokumentasi tidak melebihkan kontribusi AI dan tetap menunjukkan bagian yang diputuskan, diperbaiki, atau diverifikasi oleh manusia.",
        "Kolaborasi dengan AI tidak dilakukan melalui copy-paste hasil. Tim menemukan beberapa risiko, seperti inferensi fakta yang hilang, sumber yang tidak valid, ketidakpastian yang tidak terdeteksi, serta masalah aksesibilitas. Hasil AI direvisi menjadi kontrak structured output, validasi sumber, fallback konservatif, dan human-review gate. Validasi yang sudah dilakukan adalah 10 skenario utama lulus + 1 negative guard lulus, disertai build, lint, fallback, secret audit, dan QA aksesibilitas dasar.",
    ]
    ai_tables = [
        "Tabel 2.1 Pemetaan alat, aktivitas, risiko, dan validasi",
        "Tabel 3.1 Kronologi kolaborasi AI",
        "Tabel 4.1 Log prompt brainstorming dan perancangan",
        "Tabel 4.2 Log prompt implementasi dan debugging",
        "Tabel 4.3 Log prompt validasi dan dokumentasi",
        "Tabel 5.1 Kontrak structured output",
        "Tabel 6.1 Kesalahan AI dan koreksi manual",
        "Tabel 7.1 Hasil 10 skenario utama",
        "Tabel 7.2 Pemeriksaan teknis lainnya",
        "Tabel 8.1 Formulir evaluasi peer testing",
        "Tabel 9.1 Pembagian tanggung jawab AI dan manusia",
    ]
    ai_figures = [
        "Gambar 3.1 Bukti prompt ChatGPT untuk analisis masalah",
        "Gambar 3.2 Bukti Codex saat implementasi",
        "Gambar 4.1 Bukti prompt desain structured output",
        "Gambar 5.1 Safe Timeline Assistant",
        "Gambar 6.1 Contoh keluaran awal dan koreksi",
        "Gambar 7.1 Bukti hasil validasi terminal",
    ]
    add_front_matter(
        doc,
        ai_summary,
        "ChatGPT, Codex, Gemini, prompt engineering, validasi AI, human-in-the-loop, structured output",
        ai_tables,
        ai_figures,
    )

    chapter(doc, "I", "PENDAHULUAN")
    heading(doc, "1.1 Tujuan Dokumentasi", 2)
    paragraph(
        doc,
        "Panduan Tugas Besar AI For Real Impact 2026 mewajibkan kelompok mendokumentasikan tools AI, prompt, hasil AI, revisi, kesalahan, validasi, dan refleksi. Dokumen ini disusun sebagai catatan proses yang dapat ditinjau, bukan sebagai narasi bahwa seluruh produk dibuat otomatis. Tujuannya adalah menunjukkan bagaimana AI membantu tim berpikir dan membangun, bagian mana yang gagal, serta bagaimana manusia tetap bertanggung jawab terhadap hasil akhir.",
    )
    heading(doc, "1.2 Ruang Lingkup", 2)
    numbered(
        doc,
        [
            "AI di dalam produk: Gemini untuk menyusun draf kronologi dari catatan yang dipilih.",
            "AI sebagai partner tim: ChatGPT untuk eksplorasi masalah dan rancangan, serta Codex untuk inspeksi repository, implementasi, debugging, pengujian, dan dokumen.",
            "Validasi manusia dan otomatis: pemeriksaan sumber, skema, fallback, build, lint, browser QA, render dokumen, dan evaluasi kritis.",
            "Batas dokumentasi: tidak mengklaim peer testing atau panggilan Gemini live yang belum memiliki bukti.",
        ],
    )
    heading(doc, "1.3 Prinsip Penggunaan AI", 2)
    bullets(
        doc,
        [
            "Problem-first: AI dipilih setelah masalah dan kebutuhan pengguna dipahami.",
            "Human accountability: keputusan produk, etika, klaim, dan penerimaan hasil tetap menjadi tanggung jawab tim.",
            "Source-grounded: keluaran timeline harus terhubung dengan catatan sumber.",
            "Fail conservatively: informasi yang hilang tidak boleh ditebak.",
            "Transparent limitations dan synthetic-data only: hasil AI ditandai sebagai draf, keterbatasan prototype dijelaskan, serta pengembangan dan demonstrasi tidak menggunakan data penyintas nyata.",
        ],
    )

    chapter(doc, "II", "ALAT AI DAN PEMBAGIAN PERAN")
    heading(doc, "2.1 AI di Dalam Produk", 2)
    paragraph(
        doc,
        "Gemini digunakan pada endpoint serverless Safe Timeline Assistant. Inputnya adalah array catatan sintetis yang dipilih pengguna. Outputnya adalah struktur JSON berisi draf peristiwa. Model tidak memiliki kewenangan untuk mengirim laporan, menyatakan kebenaran, atau mengisi fakta yang tidak tersedia. UI mempertahankan keputusan pada manusia melalui edit, terima, dan tolak.",
    )
    heading(doc, "2.2 AI sebagai Partner Kerja Tim", 2)
    paragraph(
        doc,
        "ChatGPT digunakan untuk membandingkan gagasan AmanAkses dengan panduan tugas, menguji relevansi fitur AI, dan membantu merumuskan batas etika. Codex digunakan untuk membaca codebase, merancang tipe data, mengimplementasikan fungsi timeline, menambahkan endpoint, membuat fixture, memperbaiki masalah aksesibilitas, menyusun dokumen, dan menjalankan QA. Tim tidak menerima hasil secara otomatis; setiap perubahan diperiksa melalui source code, build, lint, test, render, dan inspeksi visual.",
    )
    heading(doc, "2.3 Alat Produksi Non-AI", 2)
    paragraph(
        doc,
        "Python-docx digunakan untuk menghasilkan DOCX, Microsoft Word untuk konversi PDF dan pembaruan field, PowerPoint untuk ekspor presentasi, Vite dan TypeScript untuk build, ESLint untuk pemeriksaan statis, serta PyMuPDF untuk render PDF. Alat-alat tersebut bukan AI dan dicatat terpisah agar kontribusi AI tidak dibesar-besarkan.",
    )
    table(
        doc,
        ["Aktivitas", "Alat", "Keluaran", "Risiko", "Validasi Manusia"],
        [
            ("Analisis masalah", "ChatGPT", "Pemetaan panduan, masalah, dan kandidat fitur.", "Solusi menjadi terlalu luas.", "Bandingkan dengan panduan dan riset resmi."),
            ("Perancangan AI", "ChatGPT + Codex", "Batas AI, schema, human-review gate.", "Batas hanya menjadi disclaimer.", "Implementasikan validator dan kontrol UI."),
            ("Coding", "Codex", "Endpoint, tipe, fallback, fixture, UI.", "Kode salah atau mengubah fitur lain.", "Review diff, build, lint, browser QA."),
            ("Runtime produk", "Gemini", "Draf timeline structured JSON.", "Halusinasi atau salah konteks.", "Sumber wajib, nullable fields, edit/terima/tolak."),
            ("Dokumentasi", "Codex", "Draf proposal, dokumentasi AI, dan presentasi.", "Klaim berlebihan atau format rusak.", "Riset sumber, render halaman, placeholder bukti."),
            ("Produksi dokumen", "python-docx + Word", "DOCX editable dan PDF.", "Layout terpotong.", "Render dan inspeksi setiap halaman."),
        ],
        [2.8, 2.6, 4.0, 3.7, 4.0],
        8,
    )
    caption(doc, "Tabel 2.1 Pemetaan alat, aktivitas, risiko, dan validasi")

    chapter(doc, "III", "KRONOLOGI KOLABORASI AI", page_break=False)
    heading(doc, "3.1 Tahapan Kerja", 2)
    table(
        doc,
        ["Tahap", "Pertanyaan kepada AI", "Keputusan Tim", "Artefak"],
        [
            ("1. Orientasi", "Apakah AmanAkses sesuai soal dan masalah apa yang paling kuat?", "Fokus pada masalah akses dokumentasi, bukan sekadar UI.", "Pemetaan deliverables dan rubrik."),
            ("2. Penyempitan fitur", "Fitur AI apa yang bermakna tetapi aman dan terukur?", "Pilih timeline extraction, bukan chatbot serbaguna.", "Scope Safe Timeline Assistant."),
            ("3. Kontrak data", "Bagaimana hasil tetap dapat diaudit dan tidak mengarang?", "Field nullable, sourceNoteIds, uncertainty, requiresReview.", "Types dan JSON schema."),
            ("4. Implementasi", "Bagaimana menjaga API key server-side dan demo tetap berjalan?", "Vercel function dan fallback deterministik.", "Endpoint dan library client."),
            ("5. Validasi", "Kasus apa yang dapat membongkar asumsi salah?", "10 fixture utama dan 1 negative guard.", "Validation runner."),
            ("6. QA", "Apa yang masih bermasalah pada aksesibilitas dan packaging?", "Perbaiki lang, aria-label, dokumentasi, dan secret audit.", "Build, lint, browser QA, ZIP."),
            ("7. Dokumentasi", "Bagaimana menjelaskan bukti tanpa melebihkan hasil?", "Pisahkan implementasi, simulasi, rencana, dan placeholder.", "Proposal dan dokumentasi AI."),
        ],
        [2.4, 5.0, 5.0, 3.6],
        8,
    )
    caption(doc, "Tabel 3.1 Kronologi kolaborasi AI")
    image_marker(doc, "Screenshot prompt ChatGPT untuk analisis masalah", "Gambar 3.1 Bukti prompt ChatGPT untuk analisis masalah")
    image_marker(doc, "Screenshot Codex saat implementasi atau debugging", "Gambar 3.2 Bukti Codex saat implementasi")
    heading(doc, "3.2 Pola Kolaborasi yang Digunakan", 2)
    numbered(
        doc,
        [
            "Tim memberi konteks berupa panduan tugas, tujuan AmanAkses, kode yang sudah ada, dan batas etika.",
            "AI menghasilkan alternatif, diagnosis teknis, atau draf implementasi.",
            "Tim memeriksa kesesuaian dengan masalah, sumber, prototype, dan batas risiko.",
            "Hasil yang diterima diimplementasikan atau direvisi.",
            "Build, lint, fixture, browser QA, atau render digunakan sebagai bukti eksternal terhadap jawaban AI.",
            "Kegagalan dicatat dan digunakan untuk memperbaiki prompt, validator, atau desain.",
        ],
    )

    chapter(doc, "IV", "LOG PROMPT, HASIL, DAN REVISI", page_break=False)
    heading(doc, "4.1 Brainstorming dan Perancangan", 2)
    table(
        doc,
        ["ID", "Prompt Inti", "Hasil AI", "Keputusan/Revisi"],
        [
            ("P01", "Bandingkan ide AmanAkses dengan Panduan AI For Real Impact dan identifikasi kekuatan serta kekurangannya.", "AmanAkses relevan, tetapi peran AI belum cukup konkret.", "Tambahkan satu fitur AI yang dapat didemokan dan diuji."),
            ("P02", "Pilih fitur AI paling bermakna, aman, dan dapat diuji untuk prototype AmanAkses.", "Timeline extraction dinilai lebih tepat daripada chatbot umum.", "Safe Timeline Assistant menjadi fokus AI."),
            ("P03", "Buat prinsip desain untuk pengguna disabilitas pada konteks sensitif.", "Aksesibilitas, bahasa netral, consent, safe exit, dan pilihan jeda.", "Prinsip dipakai dalam UI dan batas dokumen."),
            ("P04", "Analisis risiko bila AI menyusun kronologi pengalaman sensitif.", "Halusinasi, hilang nuansa, bias, dan overreliance.", "Sumber, uncertainty, nullable fields, dan human review diwajibkan."),
        ],
        [1.0, 5.2, 4.0, 4.8],
        8,
    )
    caption(doc, "Tabel 4.1 Log prompt brainstorming dan perancangan")
    image_marker(doc, "Screenshot prompt desain structured output", "Gambar 4.1 Bukti prompt desain structured output")

    heading(doc, "4.2 Implementasi dan Debugging", 2)
    table(
        doc,
        ["ID", "Prompt Inti", "Hasil AI", "Keputusan/Revisi"],
        [
            ("P05", "Rancang tipe data event yang dapat melacak catatan sumber dan informasi hilang.", "TimelineCandidate dengan date/time/location nullable, sourceNoteIds, uncertainty.", "requiresReview selalu true; sumber diperiksa validator."),
            ("P06", "Buat endpoint Gemini server-side dengan structured output.", "Vercel function membaca key dari environment dan mengembalikan JSON.", "Key tidak memakai prefix VITE_; tambahkan fallback pada failure."),
            ("P07", "Buat fallback konservatif agar demo berjalan tanpa API.", "Parser label tanggal, waktu, lokasi, dan ringkasan.", "recordedAt dilarang mengisi tanggal/waktu kejadian yang hilang."),
            ("P08", "Periksa UI review agar pengguna dapat menerima atau menolak setiap event.", "Editor event dan status keputusan.", "Hanya event accepted dianggap siap masuk laporan."),
            ("P09", "Cari masalah aksesibilitas pada halaman prototype.", "Ditemukan bahasa dokumen en dan input pencarian tanpa accessible name.", "Ubah lang=id dan tambahkan aria-label."),
        ],
        [1.0, 5.2, 4.0, 4.8],
        8,
    )
    caption(doc, "Tabel 4.2 Log prompt implementasi dan debugging")

    heading(doc, "4.3 Validasi dan Dokumentasi", 2)
    table(
        doc,
        ["ID", "Prompt Inti", "Hasil AI", "Keputusan/Revisi"],
        [
            ("P10", "Susun fixture untuk field hilang, waktu perkiraan, alias, dan sumber tidak valid.", "10 skenario utama dan negative test.", "Tambahkan kasus seingatku setelah kelemahan fallback ditemukan."),
            ("P11", "Audit agar secret tidak masuk browser bundle dan source archive.", "Pencarian key pattern, nama environment, dan isi ZIP.", "Hanya .env.example diarsipkan; .env dan dist dikeluarkan."),
            ("P12", "Tulis proposal sesuai struktur soal dengan status fitur yang jujur.", "Draf dokumen akademik.", "Perkuat bukti masalah, pisahkan visi dan prototype."),
            ("P13", "Tulis dokumentasi AI yang menunjukkan kesalahan dan validasi.", "Prompt log, error log, hasil test, dan refleksi.", "Hapus klaim Claude; gunakan ChatGPT, Codex, dan Gemini."),
        ],
        [1.0, 5.2, 4.0, 4.8],
        8,
    )
    caption(doc, "Tabel 4.3 Log prompt validasi dan dokumentasi")
    heading(doc, "4.4 Catatan tentang Bukti Prompt", 2)
    paragraph(
        doc,
        "Tabel prompt menyajikan isi inti dan keputusan yang dapat diaudit. Screenshot percakapan asli perlu ditambahkan oleh tim pada marker gambar jika diwajibkan oleh dosen. Sebelum ditempel, screenshot harus diperiksa agar tidak memuat API key, data pribadi, tab sensitif, atau percakapan yang tidak relevan.",
    )

    chapter(doc, "V", "AI DI DALAM AMANAKSES", page_break=False)
    heading(doc, "5.1 Tujuan dan Batas Sistem", 2)
    paragraph(
        doc,
        "Tujuan Safe Timeline Assistant adalah mengurangi beban pengguna ketika mengubah catatan terpisah menjadi urutan peristiwa yang dapat diperiksa. Model tidak digunakan untuk klasifikasi korban, deteksi kebohongan, prediksi risiko, diagnosis, identifikasi pelaku, atau rekomendasi hukum. Pembatasan tersebut bukan hanya ditulis dalam disclaimer, tetapi tercermin pada input terpilih, schema, validator, fallback, dan kontrol UI.",
    )
    heading(doc, "5.2 System Instruction", 2)
    numbered(
        doc,
        [
            "Gunakan hanya fakta yang tertulis eksplisit dalam catatan.",
            "Jangan menambah fakta, menyimpulkan niat, menilai kebenaran, menyalahkan pihak, membuat diagnosis, atau memberi nasihat hukum.",
            "Jika tanggal, waktu, atau lokasi tidak jelas, gunakan null dan tandai ketidakpastian.",
            "Gunakan bahasa Indonesia yang netral, singkat, tidak grafis, dan tidak menghakimi.",
            "Setiap event wajib mencantumkan sourceNoteIds yang benar.",
            "uncertainty hanya boleh explicit, ambiguous, atau missing.",
            "requiresReview selalu true karena semua keluaran adalah draf.",
        ],
    )
    heading(doc, "5.3 Structured-Output Contract", 2)
    table(
        doc,
        ["Field", "Tipe", "Aturan"],
        [
            ("id", "string", "ID event unik untuk kebutuhan UI."),
            ("title", "string", "Judul netral dan singkat."),
            ("date", "string atau null", "Hanya diisi jika terdapat tanggal kejadian yang eksplisit."),
            ("time", "string atau null", "Hanya diisi jika waktu eksplisit; format dinormalisasi."),
            ("location", "string atau null", "Tidak boleh ditebak."),
            ("summary", "string", "Ringkasan netral yang didukung sumber."),
            ("sourceNoteIds", "array string", "Minimal satu ID dan seluruh ID harus terdapat pada input."),
            ("uncertainty", "enum", "explicit, ambiguous, atau missing."),
            ("requiresReview", "boolean", "Harus true."),
        ],
        [3.0, 3.5, 8.2],
        9,
    )
    caption(doc, "Tabel 5.1 Kontrak structured output")
    heading(doc, "5.4 Alur Data dan Data Boundary", 2)
    numbered(
        doc,
        [
            "Browser menyimpan pilihan catatan sintetis pada state antarmuka.",
            "Hanya catatan terpilih dikirim ke POST /api/timeline.",
            "Fungsi serverless membaca GEMINI_API_KEY dari environment server.",
            "Gemini mengembalikan structured JSON yang kemudian divalidasi.",
            "Jika API gagal, client menggunakan fallback deterministik lokal.",
            "UI menampilkan sumber dan editor; pengguna menentukan status setiap event.",
            "Prototype tidak mengirim laporan ke pihak eksternal.",
        ],
    )
    image_marker(doc, "Gambar web /app/kronologi", "Gambar 5.1 Safe Timeline Assistant")
    heading(doc, "5.5 Human-Review Gate", 2)
    paragraph(
        doc,
        "Setiap event dimulai dalam status pending. Pengguna dapat membuka sumber, mengubah seluruh field, lalu memilih accepted atau rejected. Event pending atau rejected tidak dianggap siap masuk laporan. Mekanisme ini menghindari pola di mana disclaimer muncul tetapi keluaran AI langsung dipakai tanpa keputusan eksplisit.",
    )

    chapter(doc, "VI", "CONTOH HASIL, KESALAHAN, DAN KOREKSI", page_break=False)
    heading(doc, "6.1 Contoh Keluaran yang Diharapkan", 2)
    paragraph(
        doc,
        'Input sintetis: "Tanggal: 21 Mei 2026. Aku berbicara dengan pendamping tepercaya. Waktu dan lokasi tidak dicatat." Keluaran yang aman mempertahankan tanggal 21 Mei 2026, membuat time dan location bernilai null, menyertakan ID catatan sumber, memberi uncertainty sesuai kondisi field, dan menetapkan requiresReview=true.',
    )
    heading(doc, "6.2 Contoh Keluaran yang Tidak Diterima", 2)
    bullets(
        doc,
        [
            "Mengisi waktu berdasarkan recordedAt catatan meskipun waktu kejadian tidak ditulis.",
            "Mengubah percakapan menjadi kesimpulan bahwa suatu tindak pidana telah terjadi.",
            "Mengidentifikasi alias atau menebak hubungan antarorang.",
            "Mencantumkan sourceNoteIds yang tidak terdapat pada input.",
            "Menghasilkan event tanpa status review atau menganggap hasil otomatis siap dilaporkan.",
        ],
    )
    image_marker(doc, "Screenshot contoh keluaran awal dan hasil setelah koreksi", "Gambar 6.1 Contoh keluaran awal dan koreksi")
    heading(doc, "6.3 Kesalahan yang Ditemukan", 2)
    table(
        doc,
        ["Temuan", "Mengapa Berbahaya", "Koreksi Manual dan Sistem"],
        [
            ("Frasa seingatku tidak dikenali fallback.", "Informasi tidak pasti tampak eksplisit.", "Perluas deteksi ambiguitas dan tambahkan fixture."),
            ("Model dapat mencantumkan sumber yang tidak ada.", "Event tidak dapat diaudit.", "Validator menolak event tanpa ID sumber valid."),
            ("recordedAt berpotensi dianggap waktu kejadian.", "Sistem menambahkan fakta yang tidak ditulis.", "Pisahkan metadata pencatatan dari field kejadian."),
            ("Structured output terlihat meyakinkan.", "Bentuk rapi dapat disalahartikan sebagai isi yang benar.", "Tampilkan label draf, sumber, uncertainty, dan review."),
            ("Endpoint tidak tersedia pada vite dev.", "Demo terlihat rusak ketika serverless function tidak berjalan.", "Aktifkan fallback dan jelaskan mode yang digunakan."),
            ("Input pencarian tidak memiliki nama aksesibel.", "Screen reader tidak mengetahui fungsi kontrol.", "Tambahkan aria-label."),
            ("Bahasa HTML masih en.", "Pelafalan pembaca layar dapat salah.", "Ubah atribut lang menjadi id."),
        ],
        [4.4, 4.5, 5.8],
        8,
    )
    caption(doc, "Tabel 6.1 Kesalahan AI dan koreksi manual")
    heading(doc, "6.4 Revisi Prompt dan Guardrail", 2)
    paragraph(
        doc,
        "Prompt awal yang hanya meminta model menyusun kronologi belum cukup. Revisi dilakukan dengan memasukkan daftar larangan, nilai null, sumber wajib, enum ketidakpastian, dan bahasa netral. Namun prompt tetap diperlakukan sebagai satu lapisan. Validator, fallback konservatif, UI review, dan pengujian diperlukan karena model dapat mengabaikan atau salah menafsirkan instruksi.",
    )

    chapter(doc, "VII", "VALIDASI HASIL AI", page_break=False)
    heading(doc, "7.1 Sepuluh Skenario Utama", 2)
    table(
        doc,
        ["No.", "Skenario", "Ekspektasi", "Hasil"],
        [
            ("1", "Semua field eksplisit", "Tanggal, waktu, dan lokasi diambil.", "Lulus"),
            ("2", "Tanggal hilang", "date tetap null.", "Lulus"),
            ("3", "Waktu sekitar", "uncertainty menandai ambiguous.", "Lulus"),
            ("4", "Waktu hilang", "time null dan recordedAt tidak dipakai.", "Lulus"),
            ("5", "Lokasi hilang", "location tetap null.", "Lulus"),
            ("6", "Format jam 08.05", "Dinormalisasi menjadi 08:05.", "Lulus"),
            ("7", "Dua catatan", "Dua event dengan sumber masing-masing.", "Lulus"),
            ("8", "Alias R", "Alias dipertahankan tanpa identifikasi.", "Lulus"),
            ("9", "Frasa seingatku", "uncertainty menandai ambiguous.", "Lulus"),
            ("10", "Catatan non-kasus", "Ringkasan netral tanpa kesimpulan hukum.", "Lulus"),
        ],
        [1.2, 4.0, 7.5, 2.0],
        8,
    )
    caption(doc, "Tabel 7.1 Hasil 10 skenario utama: 10/10 lulus")
    heading(doc, "7.2 Negative Guard", 2)
    paragraph(
        doc,
        "Satu negative guard membuat respons dengan sourceNoteIds yang tidak terdapat pada input. Validator harus menolak seluruh respons agar event tidak masuk UI sebagai hasil yang dapat dipercaya. Pengujian ini lulus. Hasil keseluruhan ditulis sebagai 10 skenario utama lulus + 1 negative guard lulus, bukan 11/11 fixture, agar perbedaan fungsi pengujian tetap jelas.",
    )
    heading(doc, "7.3 Pemeriksaan Teknis Lain", 2)
    table(
        doc,
        ["Pemeriksaan", "Bukti", "Hasil"],
        [
            ("Build", "tsc -b dan vite build selesai tanpa error.", "Lulus"),
            ("Lint", "eslint . selesai tanpa error.", "Lulus"),
            ("Fallback", "Request failure menghasilkan mode fallback.", "Lulus"),
            ("Secret audit", "Tidak ada API-key-like string atau GEMINI_API_KEY pada bundle browser.", "Lulus"),
            ("Source archive", "Tidak memuat node_modules, dist, log, atau .env.", "Lulus"),
            ("Browser QA", "Generate, edit, accept/reject, sumber, dan banner fallback diperiksa.", "Lulus"),
            ("Aksesibilitas dasar", "lang=id, aria-label pencarian, keyboard flow, dan safe exit diperiksa.", "Diperiksa"),
            ("Gemini live", "[[ISI_BUKTI_GEMINI_LIVE_SETELAH_KEY_DIPASANG]]", "Belum diklaim"),
        ],
        [3.2, 8.2, 3.3],
        9,
    )
    caption(doc, "Tabel 7.2 Pemeriksaan teknis lainnya")
    image_marker(doc, "Screenshot terminal pnpm validate:timeline, pnpm build, dan pnpm lint", "Gambar 7.1 Bukti hasil validasi terminal")
    heading(doc, "7.4 Batas Validasi", 2)
    paragraph(
        doc,
        "Fixture sintetis menguji aturan yang sudah diketahui, bukan seluruh variasi bahasa atau konteks dunia nyata. Pengujian belum mengevaluasi trauma impact, bias terhadap ragam bahasa, atau keberhasilan pada catatan panjang. Karena itu, hasil teknis tidak boleh dipakai untuk menyatakan produk siap digunakan pada kasus nyata.",
    )

    chapter(doc, "VIII", "RENCANA PEER TESTING DAN EVALUASI", page_break=False)
    heading(doc, "8.1 Tujuan", 2)
    paragraph(
        doc,
        "Peer testing bertujuan menilai apakah pengguna memahami alur, label draf AI, sumber, ketidakpastian, keputusan accept/reject, consent, dan batas prototype. Pengujian dilakukan dengan data sintetis serta tanpa meminta peserta menceritakan pengalaman pribadi.",
    )
    heading(doc, "8.2 Skenario Tugas", 2)
    numbered(
        doc,
        [
            "Temukan halaman Pahami Kekerasan dan kembali ke dashboard.",
            "Pilih tiga catatan sintetis pada Safe Timeline Assistant.",
            "Susun draf timeline dan identifikasi mode Gemini atau fallback.",
            "Buka sumber satu event dan sebutkan informasi yang tidak pasti.",
            "Edit satu field, terima satu event, dan tolak satu event.",
            "Jelaskan event mana yang dapat masuk laporan awal.",
            "Aktifkan satu preferensi aksesibilitas dan gunakan Keluar Cepat.",
        ],
    )
    heading(doc, "8.3 Formulir Evaluasi", 2)
    table(
        doc,
        ["Pernyataan", "1", "2", "3", "4", "5", "Catatan"],
        [
            ("Saya memahami tujuan AmanAkses.", "", "", "", "", "", ""),
            ("Saya mengetahui bahwa hasil AI adalah draf.", "", "", "", "", "", ""),
            ("Saya dapat menemukan catatan sumber.", "", "", "", "", "", ""),
            ("Saya memahami perbedaan pending, diterima, dan ditolak.", "", "", "", "", "", ""),
            ("Saya memahami data apa yang akan masuk laporan.", "", "", "", "", "", ""),
            ("Navigasi dan teks mudah digunakan.", "", "", "", "", "", ""),
            ("Saya merasa dapat berhenti atau keluar dari alur.", "", "", "", "", "", ""),
        ],
        [7.1, 0.8, 0.8, 0.8, 0.8, 0.8, 3.1],
        8,
    )
    caption(doc, "Tabel 8.1 Formulir evaluasi peer testing; 1 sangat tidak setuju, 5 sangat setuju")
    heading(doc, "8.4 Hasil dan Revisi", 2)
    bullets(
        doc,
        [
            "Peserta dan tanggal: [[ISI_PESERTA_DAN_TANGGAL_UJI_PEER]].",
            "Ringkasan skor: [[ISI_RINGKASAN_SKOR_UJI_PEER]].",
            "Masalah yang paling sering ditemukan: [[ISI_MASALAH_UTAMA_UJI_PEER]].",
            "Perubahan yang dilakukan: [[ISI_REVISI_SETELAH_UJI_PEER]].",
            "Masalah yang belum diselesaikan: [[ISI_SISA_MASALAH_UJI_PEER]].",
        ],
    )

    chapter(doc, "IX", "REFLEKSI KRITIS DAN TANGGUNG JAWAB", page_break=False)
    heading(doc, "9.1 Manfaat AI bagi Kelompok", 2)
    paragraph(
        doc,
        "AI mempercepat eksplorasi alternatif, membantu mengurai dokumen tugas, menyusun tipe data, membuat fixture, dan menemukan inkonsistensi. Manfaat terbesar bukan terletak pada kemampuan menghasilkan teks atau kode dalam jumlah besar, tetapi pada percepatan siklus pertanyaan, implementasi, pemeriksaan, dan revisi.",
    )
    heading(doc, "9.2 Keterbatasan dan Risiko", 2)
    bullets(
        doc,
        [
            "Hallucination: model dapat menambah detail atau sumber yang tidak tersedia.",
            "Automation bias: keluaran yang rapi dapat diterima manusia tanpa pemeriksaan.",
            "Bias bahasa: model dapat lebih baik memahami pola bahasa tertentu dan mengabaikan ekspresi lokal atau disabilitas.",
            "Privacy leakage: konteks sensitif dapat terkirim atau tercatat jika batas data tidak dirancang.",
            "Loss of nuance: ringkasan netral dapat menghapus makna penting bagi pengguna.",
            "Dependency: layanan eksternal dapat gagal, berubah, atau menimbulkan biaya.",
        ],
    )
    heading(doc, "9.3 Kesalahan AI Terbesar", 2)
    paragraph(
        doc,
        "Kesalahan paling berbahaya bukan bug sintaks, melainkan kecenderungan membuat sistem tampak mampu menentukan kasus. Tim mengoreksi arah tersebut dengan membatasi AI pada penyusunan draf. Pada tingkat teknis, kasus seingatku menunjukkan bahwa aturan sederhana dapat salah menggolongkan ketidakpastian. Temuan ini memperkuat kesimpulan bahwa prompt dan schema tidak cukup tanpa pengujian dan tinjauan manusia.",
    )
    heading(doc, "9.4 Tanggung Jawab AI dan Manusia", 2)
    table(
        doc,
        ["Aktivitas", "AI Membantu", "Manusia Bertanggung Jawab"],
        [
            ("Analisis", "Menghasilkan alternatif dan pertanyaan.", "Memilih masalah, memeriksa sumber, dan menentukan prioritas."),
            ("Desain", "Membantu merancang flow dan schema.", "Menetapkan batas etika, aksesibilitas, dan consent."),
            ("Coding", "Menghasilkan atau memperbaiki kode.", "Review, test, keamanan, dan penerimaan perubahan."),
            ("Runtime", "Menyusun draf timeline.", "Memilih input, mengedit, menerima, menolak, dan menentukan penggunaan."),
            ("Dokumentasi", "Membantu menyusun draf.", "Menjamin kebenaran klaim, sumber, dan bukti."),
        ],
        [3.2, 5.1, 6.1],
        9,
    )
    caption(doc, "Tabel 9.1 Pembagian tanggung jawab AI dan manusia")
    heading(doc, "9.5 Kesimpulan Refleksi", 2)
    paragraph(
        doc,
        "Pengalaman pengembangan AmanAkses menunjukkan bahwa penggunaan AI yang bertanggung jawab membutuhkan pembatasan teknis, desain interaksi, validasi, dan keberanian untuk tidak mengklaim lebih dari bukti yang tersedia. AI membantu tim menata pekerjaan dan membantu pengguna menata catatan. Namun keputusan mengenai masalah, produk, kebenaran, dan tindak lanjut tetap merupakan tanggung jawab manusia.",
    )
    heading(doc, "DAFTAR REFERENSI", 1)
    add_reference_list(
        doc,
        [
            "Google AI for Developers. (2026). Gemini API documentation and structured output. https://ai.google.dev/gemini-api/docs/",
            "OpenAI. (2026). ChatGPT and Codex product documentation. https://openai.com/",
            "World Wide Web Consortium. (2023). Web Content Accessibility Guidelines (WCAG) 2.2. https://www.w3.org/TR/WCAG22/",
            "Tim Pengampu. (2026). Panduan Tugas Besar AI For Real Impact 2026: Solving Real Problems with AI Collaboration.",
            "Repository AmanAkses. Source code, fixture validasi, dokumentasi setup, dan hasil pengujian lokal. Juni 2026.",
        ],
    )

    path = EDITABLE / "Dokumentasi_AI_AmanAkses_Akademik.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    print(build_proposal())
    print(build_ai_doc())
