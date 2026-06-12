from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
EDITABLE = ROOT / "supporting" / "editable"
EDITABLE.mkdir(parents=True, exist_ok=True)
TIMELINE_SCREENSHOT = ROOT / "supporting" / "screenshots" / "timeline-review-desktop.png"
HERO = ROOT / "demo-web" / "src" / "assets" / "amanaakses-hero.png"

TEAL = "0F766E"
TEAL_DARK = "115E59"
VIOLET = "6D28D9"
INK = "172033"
MUTED = "5F6B7A"
LIGHT_TEAL = "EAF8F5"
LIGHT_VIOLET = "F4F0FF"
LIGHT_GRAY = "F4F6F8"
WHITE = "FFFFFF"
GOLD = "A16207"
RED = "B42318"


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_run_font(run, name="Calibri", size=11, color=INK, bold=False, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    run.bold = bold
    run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_row_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Halaman ")
    set_run_font(run, size=9, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def configure_document(doc: Document, title: str, preset="narrative"):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_after = Pt(7 if preset == "narrative" else 5)
    normal.paragraph_format.line_spacing = 1.24 if preset == "narrative" else 1.16
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if preset == "narrative" else WD_ALIGN_PARAGRAPH.LEFT

    heading_tokens = {
        "Heading 1": (16, TEAL_DARK, 16, 8),
        "Heading 2": (13, TEAL, 11, 5),
        "Heading 3": (11.5, VIOLET, 8, 4),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.15

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(title)
    set_run_font(run, size=8.5, color=MUTED, bold=True)
    footer = section.footer
    add_page_field(footer.paragraphs[0])


def add_cover(doc, kicker, title, subtitle, document_type, accent=TEAL):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(34)
    p.paragraph_format.space_after = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(kicker.upper())
    set_run_font(run, size=10, color=accent, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    set_run_font(run, size=30, color=INK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(20)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    set_run_font(run, size=14, color=MUTED)

    if HERO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(18)
        p.add_run().add_picture(str(HERO), width=Inches(5.85))

    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2300, 7060])
    metadata = [
        ("Dokumen", document_type),
        ("Kelompok", "[[ISI_NAMA_KELOMPOK]]"),
        ("Anggota", "[[ISI_4-5_NAMA_DAN_NIM_ANGGOTA]]"),
    ]
    for row, (label, value) in zip(table.rows, metadata):
        set_cell_shading(row.cells[0], LIGHT_TEAL)
        for cell, text, bold in ((row.cells[0], label, True), (row.cells[1], value, False)):
            cell.text = ""
            run = cell.paragraphs[0].add_run(text)
            set_run_font(run, size=10, color=TEAL_DARK if bold else INK, bold=bold)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    run = p.add_run("AI For Real Impact 2026 | Data demo seluruhnya sintetis")
    set_run_font(run, size=9.5, color=MUTED, italic=True)
    doc.add_page_break()


def add_callout(doc, title, text, fill=LIGHT_TEAL, accent=TEAL_DARK):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(title + "\n")
    set_run_font(r, size=11, color=accent, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=10.2, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.keep_together = True
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.keep_together = True
        p.add_run(item)


def add_table(doc, headers, rows, widths, header_fill=LIGHT_TEAL, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_row_repeat_header(table.rows[0])
    set_row_cant_split(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, header_fill)
        cell.text = ""
        run = cell.paragraphs[0].add_run(header)
        set_run_font(run, size=font_size, color=TEAL_DARK, bold=True)
    for values in rows:
        cells = table.add_row().cells
        set_row_cant_split(table.rows[-1])
        for index, value in enumerate(values):
            set_cell_width(cells[index], widths[index])
            set_cell_margins(cells[index])
            cells[index].text = ""
            run = cells[index].paragraphs[0].add_run(str(value))
            set_run_font(run, size=font_size, color=INK)
    return table


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(7)
    run = p.add_run(text)
    set_run_font(run, size=8.5, color=MUTED, italic=True)


def build_proposal():
    doc = Document()
    configure_document(doc, "AmanAkses | Proposal Proyek", "narrative")
    add_cover(
        doc,
        "Proposal proyek",
        "AmanAkses",
        "Platform aksesibel untuk mencatat, menyusun kronologi, dan mencari bantuan dengan kendali pengguna",
        "Proposal Tugas Besar",
    )

    doc.add_heading("Ringkasan Eksekutif", level=1)
    doc.add_paragraph(
        "AmanAkses adalah prototype web/PWA yang dirancang untuk membantu penyandang disabilitas memahami pilihan bantuan, "
        "membuat catatan pribadi, menata bukti sintetis, menyusun kronologi, dan menyiapkan laporan awal tanpa kehilangan "
        "kendali atas data. Proyek ini berangkat dari masalah nyata: proses dokumentasi dan pelaporan dapat tidak aksesibel, "
        "melelahkan, serta mengharuskan seseorang menceritakan ulang pengalaman sensitif kepada banyak pihak."
    )
    doc.add_paragraph(
        "Kontribusi AI difokuskan pada Safe Timeline Assistant. Pengguna memilih sendiri catatan yang boleh diproses; AI lalu "
        "mengekstrak peristiwa ke draft terstruktur yang menampilkan sumber dan ketidakpastian. Setiap peristiwa wajib diedit, "
        "diterima, atau ditolak oleh pengguna. Sistem tidak menilai kebenaran, tidak membuat diagnosis, dan tidak memberi keputusan hukum."
    )
    add_callout(
        doc,
        "Janji utama",
        "Mengurangi beban menyusun ulang catatan tanpa mengambil alih keputusan pengguna.",
    )

    doc.add_heading("1. Identitas Kelompok", level=1)
    add_table(
        doc,
        ["Elemen", "Isi"],
        [
            ("Nama kelompok", "[[ISI_NAMA_KELOMPOK]]"),
            ("Anggota 1", "[[ISI_NAMA_ANGGOTA_1]] | [[ISI_NIM_ANGGOTA_1]]"),
            ("Anggota 2", "[[ISI_NAMA_ANGGOTA_2]] | [[ISI_NIM_ANGGOTA_2]]"),
            ("Anggota 3", "[[ISI_NAMA_ANGGOTA_3]] | [[ISI_NIM_ANGGOTA_3]]"),
            ("Anggota 4", "[[ISI_NAMA_ANGGOTA_4]] | [[ISI_NIM_ANGGOTA_4]]"),
            ("Anggota 5", "[[OPSIONAL_NAMA_DAN_NIM_ANGGOTA_5]]"),
            ("Mata kuliah/kelas", "[[ISI_MATA_KULIAH_DAN_KELAS]]"),
        ],
        [2500, 6860],
    )

    doc.add_heading("2. Latar Belakang Masalah", level=1)
    doc.add_paragraph(
        "Penyandang disabilitas dapat menghadapi hambatan berlapis ketika membutuhkan informasi dan pendampingan terkait "
        "kekerasan: antarmuka yang tidak kompatibel dengan pembaca layar, bahasa yang terlalu kompleks, tidak tersedianya "
        "cara mencatat melalui suara, tombol yang sulit digunakan, serta proses berbagi yang tidak menjelaskan siapa akan "
        "menerima data. Hambatan tersebut dapat membuat informasi penting tercecer dan meningkatkan beban menceritakan ulang."
    )
    doc.add_paragraph(
        "Undang-Undang Nomor 8 Tahun 2016 menegaskan hak atas aksesibilitas dan akomodasi yang layak bagi penyandang "
        "disabilitas. Undang-Undang Nomor 12 Tahun 2022 dan PP Nomor 30 Tahun 2025 juga menekankan penanganan, pelindungan, "
        "pemulihan, dan aksesibilitas dalam konteks tindak pidana kekerasan seksual. Di lingkungan perguruan tinggi, "
        "Permendikbudristek Nomor 55 Tahun 2024 memperkuat kebutuhan alur pencegahan dan penanganan yang dapat diakses."
    )
    doc.add_heading("2.1 Target pengguna", level=2)
    add_bullets(
        doc,
        [
            "Pengguna utama: mahasiswa atau masyarakat penyandang disabilitas yang membutuhkan ruang dokumentasi bertahap.",
            "Pendamping tepercaya: menerima hanya informasi yang secara eksplisit disetujui pengguna.",
            "Satgas kampus dan lembaga layanan: meninjau laporan awal yang lebih terstruktur dan mencantumkan kebutuhan akses.",
            "Tim pengembang dan peneliti: mengevaluasi penerapan AI yang terbatas, transparan, dan human-in-the-loop.",
        ],
    )
    doc.add_heading("2.2 Rumusan masalah", level=2)
    add_numbered(
        doc,
        [
            "Bagaimana menyediakan alur pencatatan dan persiapan bantuan yang aksesibel serta tidak memaksa?",
            "Bagaimana membantu pengguna menyusun kronologi tanpa AI menambah fakta atau mengambil keputusan?",
            "Bagaimana memastikan data hanya diproses dan dibagikan setelah persetujuan yang jelas?",
        ],
    )

    doc.add_heading("3. Ide Solusi", level=1)
    doc.add_paragraph(
        "AmanAkses menggabungkan edukasi easy-read, jurnal, brankas bukti simulasi, Safe Timeline Assistant, pengaturan "
        "pendamping, laporan awal, direktori layanan, aksesibilitas, dan safe exit dalam satu alur. Prototype tidak langsung "
        "mendorong pengguna membuat laporan formal. Pengguna dapat belajar, mencatat sedikit demi sedikit, berhenti, meninjau, "
        "dan memilih sendiri langkah berikutnya."
    )
    doc.add_heading("3.1 Alur pengguna", level=2)
    add_numbered(
        doc,
        [
            "Mengatur preferensi aksesibilitas dan membuka ruang pribadi.",
            "Membuat atau memilih catatan sintetis yang relevan.",
            "Meminta Safe Timeline Assistant menyusun draft peristiwa.",
            "Memeriksa sumber, tanggal, waktu, lokasi, ringkasan, dan ketidakpastian.",
            "Mengedit lalu menerima atau menolak setiap peristiwa.",
            "Meninjau laporan awal dan ruang lingkup data sebelum berbagi.",
        ],
    )
    if TIMELINE_SCREENSHOT.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(TIMELINE_SCREENSHOT), width=Inches(6.45))
        add_caption(doc, "Gambar 1. Safe Timeline Assistant dalam mode fallback lokal dengan tinjauan manusia.")

    doc.add_heading("4. Peran dan Batasan AI", level=1)
    add_table(
        doc,
        ["AI boleh", "AI tidak boleh"],
        [
            ("Mengekstrak tanggal/waktu/lokasi yang tertulis.", "Menebak informasi yang hilang."),
            ("Membuat ringkasan netral dari catatan terpilih.", "Menilai benar/salah, niat, atau kesalahan pihak."),
            ("Menghubungkan event dengan ID sumber.", "Membuat diagnosis atau keputusan hukum."),
            ("Menandai informasi eksplisit, ambigu, atau hilang.", "Mengirim atau membagikan laporan otomatis."),
        ],
        [4680, 4680],
        header_fill=LIGHT_VIOLET,
    )
    add_callout(
        doc,
        "Human review wajib",
        "Semua hasil berlabel “Draft AI, belum diverifikasi”. Hanya event yang diterima pengguna yang dapat dinyatakan siap masuk laporan.",
        fill=LIGHT_VIOLET,
        accent=VIOLET,
    )

    doc.add_page_break()
    doc.add_heading("5. Fitur Utama Prototype", level=1)
    add_table(
        doc,
        ["Fitur", "Nilai bagi pengguna", "Status prototype"],
        [
            ("Profil aksesibilitas", "Ukuran teks, kontras, easy-read, kontrol besar, reduced motion.", "Interaktif"),
            ("Jurnal Aman", "Pencatatan teks/suara bertahap.", "UI simulasi"),
            ("Brankas Bukti", "Metadata, tag, hash, dan pilihan bukti.", "UI simulasi"),
            ("Safe Timeline Assistant", "Draft kronologi dengan sumber dan review.", "Interaktif + Gemini/fallback"),
            ("Pendamping & consent", "Scope akses, durasi, dan revoke.", "UI simulasi"),
            ("Laporan awal", "Ringkasan untuk ditinjau manusia.", "Preview simulasi"),
            ("Safe exit", "Mengalihkan ke halaman netral.", "Interaktif"),
        ],
        [2000, 4700, 2660],
        font_size=8.7,
    )

    doc.add_heading("6. Teknologi dan Arsitektur", level=1)
    add_bullets(
        doc,
        [
            "Frontend: React 19, TypeScript, Vite, Tailwind CSS, React Router, dan Motion.",
            "Endpoint AI: Vercel serverless function pada POST /api/timeline.",
            "Model: Gemini Flash yang dapat dikonfigurasi melalui GEMINI_MODEL.",
            "Kontrak output: JSON terstruktur dengan sourceNoteIds, date, time, location, neutralSummary, uncertainty, dan requiresReview.",
            "Fallback: ekstraksi deterministik konservatif ketika API key atau layanan Gemini tidak tersedia.",
            "Deployment: Vercel; GEMINI_API_KEY hanya tersedia di server dan tidak memakai awalan VITE_.",
        ],
    )
    add_callout(
        doc,
        "Batas data",
        "Prototype dan seluruh pengujian memakai data sintetis. Catatan nyata tidak boleh dimasukkan ke deployment demo kelas.",
        fill="FFF7E8",
        accent=GOLD,
    )

    doc.add_heading("7. Validasi dan Evaluasi", level=1)
    doc.add_paragraph(
        "Validasi teknis saat penyusunan paket menunjukkan build dan lint lulus. Sepuluh fixture sintetis juga lulus, "
        "mencakup informasi lengkap, tanggal hilang, waktu perkiraan, waktu/lokasi hilang, format jam bertitik, dua catatan, "
        "alias, ingatan tidak pasti, dan catatan non-kasus. Respons dengan sourceNoteIds yang tidak dikenal ditolak."
    )
    add_table(
        doc,
        ["Pemeriksaan", "Kriteria", "Hasil"],
        [
            ("Build TypeScript/Vite", "Tidak ada error kompilasi.", "Lulus"),
            ("ESLint", "Tidak ada error lint.", "Lulus"),
            ("10 fixture timeline", "Semua output sesuai harapan.", "10/10 lulus"),
            ("Sumber event", "Setiap event memiliki sumber valid.", "Lulus"),
            ("Fallback", "Demo tetap berjalan tanpa endpoint/API.", "Lulus"),
            ("Aksesibilitas dasar", "Bahasa dokumen, label form, heading, alt text.", "Lulus setelah revisi"),
            ("Uji pengguna", "3–5 rekan menyelesaikan alur.", "[[ISI_HASIL_UJI_PEER]]"),
        ],
        [2600, 4500, 2260],
        font_size=8.8,
    )

    doc.add_heading("8. Keamanan, Etika, dan Aksesibilitas", level=1)
    add_bullets(
        doc,
        [
            "Privacy by design: minimisasi input dan kunci API hanya di server.",
            "Consent by design: pemilihan catatan dan penerimaan event dilakukan eksplisit.",
            "Human-in-the-loop: hasil AI tidak pernah otomatis menjadi laporan final.",
            "Trauma-informed flow: bahasa netral, tidak grafis, dapat dijeda, dan tidak memaksa.",
            "Accessibility by design: semantik, label form, navigasi keyboard, kontras, easy-read, kontrol besar, dan reduced motion.",
            "No silent logging: implementasi produksi tidak boleh mencatat isi catatan sensitif.",
        ],
    )
    doc.add_paragraph(
        "Sebelum digunakan di dunia nyata, AmanAkses memerlukan threat modeling, autentikasi, enkripsi end-to-end yang "
        "terverifikasi, kebijakan retensi, audit keamanan independen, serta co-design bersama organisasi disabilitas dan lembaga layanan."
    )

    doc.add_heading("9. Dampak dan Pengembangan Lanjutan", level=1)
    add_bullets(
        doc,
        [
            "Mengurangi pekerjaan manual ketika pengguna menata catatan menjadi urutan yang dapat diperiksa.",
            "Membuat keterbatasan AI terlihat melalui sumber dan status ketidakpastian.",
            "Mendukung komunikasi kebutuhan akses kepada pendamping.",
            "Menjadi dasar proposal GEMASTIK, penelitian HCI/AI safety, atau pengabdian masyarakat setelah validasi etis.",
        ],
    )
    doc.add_heading("9.1 Tahap lanjutan", level=2)
    add_numbered(
        doc,
        [
            "Co-design dan usability testing dengan komunitas disabilitas serta tenaga pendamping.",
            "Backend aman, autentikasi, penyimpanan terenkripsi, consent ledger, dan audit log.",
            "Evaluasi model menggunakan dataset sintetis yang lebih beragam dan red-team prompt.",
            "Pilot terbatas bersama lembaga yang memiliki SOP perlindungan data dan respons kasus.",
        ],
    )

    doc.add_heading("Referensi", level=1)
    refs = [
        "BPK RI. Undang-Undang Nomor 8 Tahun 2016 tentang Penyandang Disabilitas. https://peraturan.bpk.go.id/Details/37251/uu-no-8-tahun-2016",
        "BPK RI. Undang-Undang Nomor 12 Tahun 2022 tentang Tindak Pidana Kekerasan Seksual. https://peraturan.bpk.go.id/Details/207944/uu-no-12-tahun-2022",
        "BPK RI. Permendikbudristek Nomor 55 Tahun 2024 tentang Pencegahan dan Penanganan Kekerasan di Lingkungan Perguruan Tinggi.",
        "BPK RI. PP Nomor 30 Tahun 2025 tentang Pencegahan TPKS serta Penanganan, Pelindungan, dan Pemulihan Korban.",
        "W3C. Web Content Accessibility Guidelines (WCAG) 2.2. https://www.w3.org/TR/WCAG22/",
        "Google AI for Developers. Gemini API: structured outputs dan model documentation. https://ai.google.dev/gemini-api/docs/",
        "Panduan Tugas Besar AI For Real Impact 2026.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(-0.22)
        p.paragraph_format.left_indent = Inches(0.22)
        p.add_run(ref)

    path = EDITABLE / "Proposal_AmanAkses.docx"
    doc.save(path)
    return path


def build_ai_documentation():
    doc = Document()
    configure_document(doc, "AmanAkses | Dokumentasi Penggunaan AI", "compact")
    add_cover(
        doc,
        "Dokumentasi AI",
        "AmanAkses",
        "Bukti penggunaan, validasi, revisi, dan refleksi AI yang bertanggung jawab",
        "Dokumentasi Penggunaan AI",
        accent=VIOLET,
    )

    doc.add_heading("Ringkasan", level=1)
    doc.add_paragraph(
        "Dokumen ini membedakan dua bentuk penggunaan AI: AI sebagai fitur di dalam AmanAkses dan AI sebagai partner kerja "
        "tim selama problem framing, desain, coding, debugging, dan penyusunan artefak. Semua hasil AI diperiksa, dibatasi, "
        "dan direvisi. Bukti yang belum tersedia, seperti uji peer dan panggilan Gemini live, diberi placeholder eksplisit."
    )
    add_table(
        doc,
        ["Kategori", "Tujuan", "Keluaran"],
        [
            ("AI di dalam produk", "Menyusun draft timeline dari catatan terpilih.", "Event JSON dengan sumber dan uncertainty."),
            ("AI sebagai partner tim", "Membantu eksplorasi, implementasi, debugging, dan dokumentasi.", "Ide, kode awal, revisi, pengujian, dokumen."),
        ],
        [2600, 3400, 3360],
        header_fill=LIGHT_VIOLET,
    )

    doc.add_heading("1. AI Tools yang Digunakan", level=1)
    add_table(
        doc,
        ["Aktivitas", "Tool", "Penggunaan", "Validasi manusia"],
        [
            ("Analisis masalah", "ChatGPT/Codex", "Membandingkan panduan tugas dengan AmanAkses.", "Membaca ulang panduan dan proposal."),
            ("Perancangan fitur", "ChatGPT/Codex", "Memilih Safe Timeline Assistant dan batas etik.", "Menyempitkan fitur menjadi satu alur."),
            ("Coding", "Codex", "Membuat kontrak data, UI review, endpoint, fallback, fixture.", "Build, lint, browser QA, inspeksi kode."),
            ("Model runtime", "Gemini API", "Structured extraction pada endpoint server.", "Schema validation dan human review."),
            ("Desain dokumen", "Codex + python-docx", "Menyusun proposal dan dokumentasi.", "Render PDF dan inspeksi setiap halaman."),
            ("Presentasi", "Codex artifact tooling", "Membuat deck editable dan PDF.", "Render slide dan inspeksi visual."),
        ],
        [1500, 1550, 3300, 3010],
        font_size=8.2,
        header_fill=LIGHT_VIOLET,
    )

    doc.add_heading("2. Log Prompt dan Revisi", level=1)
    prompt_rows = [
        (
            "P1 — Penentuan fokus",
            "“Bandingkan AmanAkses dengan rubric AI For Real Impact. Pilih satu fitur AI yang paling bermakna, aman, dan dapat diuji.”",
            "Memilih timeline extraction, bukan chatbot umum.",
            "Menambahkan syarat source reference dan human review.",
        ),
        (
            "P2 — Kontrak output",
            "“Rancang JSON untuk event timeline yang tidak boleh mengarang fakta dan dapat melacak catatan sumber.”",
            "Skema event, date/time/location nullable, sourceNoteIds.",
            "requiresReview dipaksa true dan sumber divalidasi.",
        ),
        (
            "P3 — Guardrail",
            "“Tulis system instruction yang melarang vonis, diagnosis, inferensi niat, dan fakta baru.”",
            "Daftar larangan dan bahasa netral.",
            "Menambah null untuk informasi tidak jelas.",
        ),
        (
            "P4 — Fallback",
            "“Buat fallback lokal konservatif agar demo tetap berjalan tanpa API.”",
            "Ekstraksi label tanggal, waktu, dan lokasi.",
            "Tidak memakai recordedAt untuk mengisi event yang hilang.",
        ),
        (
            "P5 — Pengujian",
            "“Susun skenario sintetis untuk missing fields, ambiguity, alias, dan invalid source.”",
            "Sepuluh fixture dan satu negative test.",
            "Menambahkan kasus ‘seingatku’ setelah tes awal gagal.",
        ),
    ]
    add_table(
        doc,
        ["ID", "Prompt ringkas", "Hasil AI", "Revisi/validasi"],
        prompt_rows,
        [1350, 3300, 2250, 2460],
        font_size=7.8,
        header_fill=LIGHT_VIOLET,
    )
    doc.add_paragraph(
        "Catatan bukti: dokumentasi final perlu ditambah screenshot percakapan/prompt asli tim apabila dosen mensyaratkan bukti visual. "
        "Placeholder: [[TEMPEL_SCREENSHOT_PROMPT_TIM_ASLI]]."
    )

    doc.add_heading("3. System Instruction Safe Timeline Assistant", level=1)
    add_callout(
        doc,
        "Tujuan",
        "Mengekstrak peristiwa dari catatan sintetis terpilih ke draft terstruktur tanpa menambah fakta.",
        fill=LIGHT_VIOLET,
        accent=VIOLET,
    )
    system_prompt = [
        "Gunakan hanya fakta yang tertulis eksplisit dalam catatan.",
        "Jangan menambah fakta, menyimpulkan niat, menilai kebenaran, menyalahkan pihak, membuat diagnosis, atau memberi klasifikasi/nasihat hukum.",
        "Jika tanggal, waktu, atau lokasi tidak jelas, isi dengan null.",
        "Gunakan bahasa Indonesia yang netral, singkat, dan tidak grafis.",
        "Setiap event wajib mencantumkan sourceNoteIds yang benar.",
        "uncertainty harus explicit, ambiguous, atau missing.",
        "requiresReview selalu true.",
        "Semua hasil adalah draft untuk ditinjau manusia.",
    ]
    add_numbered(doc, system_prompt)

    doc.add_page_break()
    doc.add_heading("4. Kontrak Structured Output", level=1)
    add_table(
        doc,
        ["Field", "Tipe", "Aturan"],
        [
            ("id", "string", "ID event; dapat dinormalisasi aplikasi."),
            ("sourceNoteIds", "string[]", "Wajib minimal satu dan harus termasuk input."),
            ("date", "string | null", "Null jika tidak eksplisit."),
            ("time", "string | null", "Null jika tidak eksplisit."),
            ("location", "string | null", "Null jika tidak eksplisit."),
            ("title", "string", "Judul netral, maksimal 120 karakter setelah sanitasi."),
            ("neutralSummary", "string", "Tidak menambah fakta; maksimal 1000 karakter."),
            ("uncertainty", "enum", "explicit, ambiguous, atau missing."),
            ("requiresReview", "true", "Dipaksa true oleh aplikasi."),
        ],
        [2200, 1900, 5260],
        font_size=8.7,
        header_fill=LIGHT_VIOLET,
    )

    doc.add_heading("5. Contoh Input dan Output", level=1)
    doc.add_heading("5.1 Input sintetis", level=2)
    add_callout(
        doc,
        "note-campus",
        "Tanggal: 20 Mei 2026. Waktu: 13:30. Lokasi: area publik kampus. Aku mencatat sebuah situasi yang membuatku merasa tidak aman. Aku belum ingin menyimpulkan apa pun.",
        fill=LIGHT_GRAY,
        accent=INK,
    )
    doc.add_heading("5.2 Output fallback yang telah divalidasi", level=2)
    add_table(
        doc,
        ["Field", "Nilai"],
        [
            ("sourceNoteIds", '["note-campus"]'),
            ("date / time", "20 Mei 2026 / 13:30"),
            ("location", "area publik kampus"),
            ("title", "Catatan setelah kegiatan kampus"),
            ("uncertainty", "explicit"),
            ("requiresReview", "true"),
        ],
        [2600, 6760],
        font_size=9,
        header_fill=LIGHT_GRAY,
    )
    doc.add_paragraph(
        "Status Gemini live: [[ISI_BUKTI_OUTPUT_GEMINI_LIVE_SETELAH_API_KEY_DIPASANG]]. Paket ini tidak mengklaim panggilan live "
        "karena kunci API tidak disimpan di repository."
    )

    doc.add_heading("6. Kesalahan AI/Implementasi yang Ditemukan", level=1)
    add_table(
        doc,
        ["Temuan", "Risiko", "Perbaikan"],
        [
            ("Istilah “seingatku” awalnya tidak dikenali fallback.", "Informasi tidak pasti tampak eksplisit.", "Regex diperluas menjadi seingat\\w* dan fixture diulang."),
            ("Respons model dapat menyertakan source ID yang tidak ada.", "Event tidak dapat diaudit.", "Validator menolak event tanpa sumber valid."),
            ("Endpoint lokal tidak tersedia di vite dev.", "Demo tombol terlihat gagal.", "Client otomatis menggunakan fallback deterministik."),
            ("Placeholder search bukan label aksesibel.", "Pengguna screen reader tidak mendapat nama kontrol stabil.", "Menambahkan aria-label."),
            ("HTML masih lang=en.", "Pelafalan pembaca layar dapat salah.", "Mengubah document language menjadi id."),
            ("Model dapat salah walau memakai structured output.", "Ringkasan dapat menyimpang dari konteks.", "Sumber ditampilkan dan review per event diwajibkan."),
        ],
        [2800, 2900, 3660],
        font_size=8.4,
        header_fill="FFF7E8",
    )

    doc.add_heading("7. Validasi Sepuluh Skenario Sintetis", level=1)
    fixture_rows = [
        ("1", "Semua field eksplisit", "Tanggal/waktu/lokasi diekstrak.", "Lulus"),
        ("2", "Tanggal hilang", "date tetap null.", "Lulus"),
        ("3", "Waktu “sekitar”", "uncertainty=ambiguous.", "Lulus"),
        ("4", "Waktu hilang", "time=null; tidak memakai recordedAt.", "Lulus"),
        ("5", "Lokasi hilang", "location=null.", "Lulus"),
        ("6", "Jam 08.05", "Dinormalisasi menjadi 08:05.", "Lulus"),
        ("7", "Dua catatan", "Dua event dan sumber masing-masing.", "Lulus"),
        ("8", "Alias R", "Alias dipertahankan tanpa identifikasi.", "Lulus"),
        ("9", "Kata seingatku", "uncertainty=ambiguous.", "Lulus"),
        ("10", "Catatan non-kasus", "Tetap netral; tanpa kesimpulan hukum.", "Lulus"),
        ("Negatif", "Sumber dibuat-buat", "Validator menolak respons.", "Lulus"),
    ]
    add_table(
        doc,
        ["No.", "Skenario", "Ekspektasi", "Hasil"],
        fixture_rows,
        [800, 2600, 4300, 1660],
        font_size=8.2,
        header_fill=LIGHT_VIOLET,
    )
    add_callout(
        doc,
        "Hasil otomatis",
        "10/10 fixture lulus. Negative test untuk event tanpa sumber valid juga berhasil ditolak.",
        fill=LIGHT_TEAL,
        accent=TEAL_DARK,
    )

    doc.add_heading("8. Form Uji Peer 3–5 Orang", level=1)
    doc.add_paragraph(
        "Gunakan data sintetis. Penguji tidak perlu menceritakan pengalaman pribadi. Catat hasil di tabel berikut dan ringkas revisi yang dilakukan."
    )
    add_table(
        doc,
        ["Tugas", "Kriteria keberhasilan", "Catatan pengamat"],
        [
            ("Pilih dua catatan", "Penguji memahami hanya catatan terpilih yang diproses.", "[[ISI]]"),
            ("Susun timeline", "Penguji menemukan tombol dan memahami loading/fallback.", "[[ISI]]"),
            ("Temukan sumber", "Penguji dapat menunjukkan asal sebuah event.", "[[ISI]]"),
            ("Perbaiki kesalahan", "Penguji mengedit satu field dan status kembali perlu cek.", "[[ISI]]"),
            ("Terima/tolak event", "Penguji memahami perbedaan status.", "[[ISI]]"),
            ("Keluar cepat", "Penguji menemukan safe exit tanpa bantuan.", "[[ISI]]"),
        ],
        [2600, 4000, 2760],
        font_size=8.5,
    )
    add_table(
        doc,
        ["Penguji", "Waktu selesai", "Skor kemudahan 1–5", "Masalah utama"],
        [
            ("[[PEER_1]]", "[[ISI]]", "[[ISI]]", "[[ISI]]"),
            ("[[PEER_2]]", "[[ISI]]", "[[ISI]]", "[[ISI]]"),
            ("[[PEER_3]]", "[[ISI]]", "[[ISI]]", "[[ISI]]"),
            ("[[OPSIONAL_PEER_4]]", "[[ISI]]", "[[ISI]]", "[[ISI]]"),
            ("[[OPSIONAL_PEER_5]]", "[[ISI]]", "[[ISI]]", "[[ISI]]"),
        ],
        [1900, 1800, 2200, 3460],
        font_size=8.4,
        header_fill=LIGHT_GRAY,
    )

    doc.add_heading("9. Refleksi Kritis", level=1)
    doc.add_heading("9.1 Manfaat AI", level=2)
    add_bullets(
        doc,
        [
            "Mempercepat eksplorasi banyak alternatif solusi dan membantu mempersempit fokus.",
            "Membantu menghasilkan kontrak data dan test case secara konsisten.",
            "Membuat implementasi dan dokumentasi lebih cepat, sehingga tim dapat mengalokasikan waktu untuk validasi.",
        ],
    )
    doc.add_heading("9.2 Keterbatasan dan risiko", level=2)
    add_bullets(
        doc,
        [
            "Model dapat menghasilkan ringkasan yang tampak meyakinkan tetapi tidak didukung catatan.",
            "Structured output menjamin bentuk data, bukan kebenaran isinya.",
            "Bahasa netral dapat menghapus nuansa penting atau tetap terasa tidak sesuai bagi pengguna.",
            "Data sensitif dapat terekspos jika arsitektur logging, retensi, dan persetujuan tidak dirancang dengan benar.",
            "Prototype belum menggantikan co-design dengan penyandang disabilitas dan pendamping profesional.",
        ],
    )
    doc.add_heading("9.3 Kesalahan terbesar yang dihindari", level=2)
    doc.add_paragraph(
        "Kesalahan terbesar adalah menjadikan AI sebagai penentu kasus atau pemberi rekomendasi hukum. Tim memilih fungsi "
        "ekstraksi yang terbatas, dapat diaudit, dan dapat dilewati. Keberhasilan tidak diukur dari seberapa banyak teks yang "
        "dibuat AI, tetapi dari apakah pengguna dapat menemukan kesalahan dan mempertahankan kendali."
    )

    doc.add_heading("10. Checklist Sebelum Pengumpulan", level=1)
    add_bullets(
        doc,
        [
            "[[ISI_IDENTITAS_KELOMPOK_DI_SEMUA_DOKUMEN]]",
            "[[PASANG_GEMINI_API_KEY_DI_VERCEL_DAN_UJI_LIVE]]",
            "[[ISI_HASIL_UJI_PEER_DAN_REVISI]]",
            "[[TEMPEL_SCREENSHOT_PROMPT_TIM_ASLI]]",
            "[[ISI_LINK_PROTOTYPE]]",
            "[[ISI_LINK_VIDEO_DEMO]]",
            "Pastikan SourceCode.zip tidak berisi .env, node_modules, dist, atau log.",
            "Putar demo dari awal sampai safe exit sebelum merekam video.",
        ],
    )

    path = EDITABLE / "Dokumentasi_AI_AmanAkses.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    proposal = build_proposal()
    ai_doc = build_ai_documentation()
    print(proposal)
    print(ai_doc)
