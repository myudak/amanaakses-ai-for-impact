# -*- coding: utf-8 -*-
"""Build a complete Tugas Proyek AI document from the user's original DOCX.

The original cover, logo, team identity, section geometry, and general visual
language are preserved. The incomplete body is replaced with a fuller,
evidence-based report that matches the implemented AmanAkses prototype.
"""

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / "supporting" / "tubes-ai-document" / "TUBES_AI_original.docx"
OUTPUT = ROOT / "outputs" / "tugas-besar-ai" / "TUBES_AI_Lengkap.docx"

DIAGRAMS = ROOT / "supporting" / "diagrams"
SCREENSHOTS = ROOT / "supporting" / "frontend-qa"
FEEDBACK_PREVIEW = ROOT / "supporting" / "feedback-workbook-preview"

FONT = "Times New Roman"
BLACK = RGBColor(0, 0, 0)
GRAY = "E7E6E6"
LIGHT_GRAY = "F2F2F2"


def set_run_font(run, size=12, bold=False, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.color.rgb = BLACK
    run.bold = bold
    run.italic = italic
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), FONT)


def set_outline_level(paragraph, level):
    ppr = paragraph._p.get_or_add_pPr()
    outline = ppr.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        ppr.append(outline)
    outline.set(qn("w:val"), str(level))


def set_keep_with_next(paragraph, enabled=True):
    paragraph.paragraph_format.keep_with_next = enabled


def add_field(paragraph, instruction, placeholder=""):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    return run


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(12)
    normal.font.color.rgb = BLACK
    normal.element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal.element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(0)

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[style_name]
        style.font.name = FONT
        style.font.size = Pt(12)
        style.font.bold = True
        style.font.color.rgb = BLACK
        style.element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style.element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.keep_with_next = True


def truncate_after_cover(doc):
    body = doc._element.body
    start = None
    for child in list(body):
        if child.tag == qn("w:p"):
            text = "".join(node.text or "" for node in child.iter(qn("w:t"))).strip()
            if text == "KATA PENGANTAR":
                start = child
                break
    if start is None:
        raise RuntimeError("KATA PENGANTAR marker not found in source document.")

    deleting = False
    for child in list(body):
        if child is start:
            deleting = True
        if deleting and child.tag != qn("w:sectPr"):
            body.remove(child)

    # The source cover uses Heading 1 for several visual lines. Keep all direct
    # formatting, but use Normal so Word does not pull cover text into the TOC.
    for paragraph in doc.paragraphs:
        paragraph.style = doc.styles["Normal"]

    last_paragraph = doc.paragraphs[-1]
    last_paragraph.add_run().add_break(WD_BREAK.PAGE)


def add_page_numbers(doc):
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.clear()
    add_field(paragraph, "PAGE", "1")
    for run in paragraph.runs:
        set_run_font(run, size=10)


def add_body(
    doc,
    text,
    *,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_indent=True,
    bold=False,
    italic=False,
    size=12,
    before=0,
    after=0,
    keep=False,
):
    paragraph = doc.add_paragraph()
    paragraph.alignment = align
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.keep_with_next = keep
    if first_indent:
        paragraph.paragraph_format.first_line_indent = Cm(1.0)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    return paragraph


def add_center_title(doc, text, *, outline=None, page_break_before=False):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.page_break_before = page_break_before
    run = paragraph.add_run(text)
    set_run_font(run, bold=True)
    if outline is not None:
        set_outline_level(paragraph, outline)
    return paragraph


def add_chapter(doc, roman, title):
    paragraph = add_center_title(
        doc,
        f"BAB {roman}\n{title}",
        outline=0,
        page_break_before=True,
    )
    paragraph.paragraph_format.space_after = Pt(12)
    return paragraph


def add_subheading(doc, number, title):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(f"{number} {title}")
    set_run_font(run, bold=True)
    set_outline_level(paragraph, 1)
    return paragraph


def add_minor_heading(doc, title):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(title)
    set_run_font(run, bold=True)
    return paragraph


def add_bullet(doc, text, *, number=None):
    prefix = f"{number}. " if number is not None else "• "
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.left_indent = Cm(0.75)
    paragraph.paragraph_format.first_line_indent = Cm(-0.5)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(prefix + text)
    set_run_font(run)
    return paragraph


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=70, start=90, bottom=70, end=90):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def add_table(doc, headers, rows, widths=None, font_size=10):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_borders(table)
    table.autofit = False
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = ""
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, GRAY)
        set_cell_margins(cell)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.line_spacing = 1.0
        run = paragraph.add_run(header)
        set_run_font(run, size=font_size, bold=True)

    header_props = table.rows[0]._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    header_props.append(repeat)

    for row_index, row_values in enumerate(rows):
        cells = table.add_row().cells
        if row_index % 2 == 1:
            for cell in cells:
                set_cell_shading(cell, LIGHT_GRAY)
        for index, value in enumerate(row_values):
            cell = cells[index]
            cell.text = ""
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_margins(cell)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(str(value))
            set_run_font(run, size=font_size)

    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Cm(width)

    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_caption(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_run_font(run, size=10, bold=False)
    return paragraph


def add_image(doc, path, caption, width_cm=14.7):
    if not path.exists():
        add_body(
            doc,
            f"[GAMBAR TIDAK DITEMUKAN: {path.name}]",
            align=WD_ALIGN_PARAGRAPH.CENTER,
            first_indent=False,
            italic=True,
        )
        add_caption(doc, caption)
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    add_caption(doc, caption)


def add_toc(doc):
    add_center_title(doc, "DAFTAR ISI")
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    add_field(
        paragraph,
        'TOC \\o "1-2" \\h \\z \\u',
        "Klik kanan lalu pilih Update Field untuk memperbarui daftar isi.",
    )
    for run in paragraph.runs:
        set_run_font(run)
    add_body(
        doc,
        "Catatan: setelah mengubah isi atau gambar, tekan Ctrl+A lalu F9 di Microsoft Word agar daftar isi dan nomor halaman diperbarui.",
        first_indent=False,
        italic=True,
        size=10,
    )


def add_front_matter(doc):
    add_center_title(doc, "KATA PENGANTAR")
    add_body(
        doc,
        "Puji syukur ke hadirat Tuhan Yang Maha Esa karena atas rahmat-Nya laporan Tugas Proyek Kecerdasan Buatan berjudul “Sistem AmanAkses” dapat diselesaikan. Dokumen ini disusun sebagai bentuk penerapan kecerdasan buatan untuk membantu menyelesaikan masalah sosial nyata melalui pendekatan yang bertanggung jawab, dapat divalidasi, dan tetap menempatkan manusia sebagai pengambil keputusan.",
    )
    add_body(
        doc,
        "AmanAkses dirancang sebagai prototype platform digital aksesibel bagi penyandang disabilitas untuk memahami informasi, membuat catatan bertahap, menyusun kronologi, dan menyiapkan bahan awal sebelum mencari bantuan. Kecerdasan buatan pada prototype tidak digunakan untuk menentukan kebenaran suatu kejadian, membuat diagnosis, menetapkan kesalahan, atau menggantikan pendamping profesional. Perannya dibatasi pada penyusunan draf informasi dan navigasi yang selalu memerlukan tinjauan serta persetujuan pengguna.",
    )
    add_body(
        doc,
        "Penyusunan proyek ini melibatkan studi regulasi dan aksesibilitas, perancangan pengalaman pengguna, implementasi prototype, pengujian fixture sintetis, pemeriksaan build dan lint, serta dokumentasi penggunaan AI oleh tim. Seluruh data contoh yang digunakan dalam aplikasi, gambar, dan analisis feedback bersifat sintetis. Pengujian bersama pengguna nyata masih menjadi bagian dari rencana pengembangan lanjutan.",
    )
    add_body(
        doc,
        "Tim menyadari bahwa laporan dan prototype ini masih memiliki keterbatasan. Oleh karena itu, kritik dan saran sangat diharapkan untuk memperbaiki aksesibilitas, keamanan, ketepatan penggunaan AI, dan relevansi sistem bagi pengguna. Semoga proyek ini dapat menjadi langkah awal untuk pengembangan teknologi yang lebih inklusif dan berdampak.",
    )
    add_body(
        doc,
        "Semarang, Juni 2026\n\nTim Penyusun",
        align=WD_ALIGN_PARAGRAPH.RIGHT,
        first_indent=False,
    )
    doc.add_page_break()
    add_toc(doc)
    doc.add_page_break()

    add_center_title(doc, "DAFTAR TABEL")
    tables = [
        "Tabel 2.1 Persona dan kebutuhan akses",
        "Tabel 2.2 Kebutuhan fungsional prioritas",
        "Tabel 2.3 Kebutuhan non-fungsional",
        "Tabel 3.1 Pemetaan AI dan alat produksi",
        "Tabel 3.2 Log prompt dan keputusan manusia",
        "Tabel 3.3 Kesalahan AI dan koreksi",
        "Tabel 4.1 Fitur, manfaat, dan status prototype",
        "Tabel 4.2 Teknologi prototype dan visi produksi",
        "Tabel 6.1 Hasil validasi teknis",
        "Tabel 6.2 Ringkasan data feedback sintetis",
        "Tabel 6.3 Risiko dan mitigasi",
    ]
    for item in tables:
        add_body(doc, item, first_indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)

    add_center_title(doc, "DAFTAR GAMBAR")
    figures = [
        "Gambar 3.1 Arsitektur sistem AmanAkses",
        "Gambar 4.1 Alur pengguna dan kendali AI",
        "Gambar 5.1 Halaman Pahami Kekerasan",
        "Gambar 5.2 Asisten Aman dan tool terstruktur",
        "Gambar 5.3 Workspace Safe Timeline Assistant",
        "Gambar 5.4 Review peristiwa dan referensi sumber",
        "Gambar 6.1 Dashboard feedback pengguna sintetis",
    ]
    for item in figures:
        add_body(doc, item, first_indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)


def add_chapter_one(doc):
    add_chapter(doc, "I", "PENDAHULUAN")
    add_subheading(doc, "1.1", "Latar Belakang")
    paragraphs = [
        "Kekerasan seksual merupakan permasalahan serius yang membutuhkan pencegahan, penanganan, pelindungan, dan pemulihan secara terpadu. Dalam praktiknya, seseorang yang mengalami atau mengkhawatirkan suatu kejadian tidak selalu langsung siap membuat laporan formal. Pengguna mungkin terlebih dahulu perlu memahami situasi, mengingat detail, menyimpan catatan, mencari orang yang dipercaya, atau mengetahui bentuk bantuan yang tersedia. Tahap persiapan ini sering tidak terwadahi oleh sistem digital yang hanya menyediakan formulir pengaduan.",
        "Bagi penyandang disabilitas, hambatan tersebut dapat menjadi berlapis. Informasi layanan sering ditulis dalam bahasa hukum yang panjang, alur situs sulit dioperasikan dengan keyboard, label antarmuka tidak terbaca pembaca layar, pilihan komunikasi terbatas, dan prosedur dapat menuntut pengguna menceritakan pengalaman berulang kali. Pengguna dengan kebutuhan kognitif atau psikososial juga dapat memerlukan bahasa sederhana, langkah kecil, waktu jeda, dan kendali untuk menghentikan proses tanpa kehilangan pekerjaan.",
        "Komnas Perempuan dalam Siaran Pers Hari Disabilitas Internasional 2024 menyebutkan bahwa perempuan dengan disabilitas masih menghadapi diskriminasi sistemik serta risiko kekerasan berbasis gender dan disabilitas. CATAHU 2024 mencatat 105 kasus kekerasan terhadap perempuan dengan disabilitas dan 38 kasus di antaranya dilaporkan langsung ke Komnas Perempuan. Angka tersebut tidak dapat dianggap mewakili seluruh kejadian, tetapi menunjukkan urgensi mekanisme informasi, dokumentasi, dan rujukan yang lebih inklusif.",
        "Kerangka hukum Indonesia telah menegaskan hak aksesibilitas dan akomodasi yang layak. Undang-Undang Nomor 8 Tahun 2016 mengatur hak penyandang disabilitas; Undang-Undang Nomor 12 Tahun 2022 mengatur pencegahan dan penanganan tindak pidana kekerasan seksual; Permendikbudristek Nomor 55 Tahun 2024 mengatur pencegahan dan penanganan kekerasan di lingkungan perguruan tinggi; serta Peraturan Pemerintah Nomor 30 Tahun 2025 menegaskan aksesibilitas dalam pencegahan, penanganan, pelindungan, dan pemulihan korban.",
        "Perkembangan Large Language Model memberi peluang untuk membantu pekerjaan administratif-kognitif, misalnya merapikan catatan menjadi urutan waktu atau menjelaskan menu dengan bahasa yang lebih sederhana. Namun, teknologi yang sama dapat mengarang detail, menghilangkan nuansa, atau membuat kesimpulan yang tidak bersumber. Oleh karena itu, pemanfaatan AI pada konteks sensitif harus dibatasi, dapat diaudit, memiliki jalur tanpa AI, dan selalu mewajibkan tinjauan manusia.",
        "Berdasarkan masalah tersebut, tim mengembangkan AmanAkses, yaitu prototype platform digital aksesibel yang membantu pengguna memahami pilihan, membuat jurnal, menata bukti sintetis, menyusun draf kronologi, memilih pendamping, dan meninjau laporan awal. AmanAkses bukan layanan darurat, bukan konselor, bukan alat diagnosis, dan bukan sistem penentu kebenaran. Produk diposisikan sebagai ruang persiapan yang menempatkan kendali pada pengguna.",
    ]
    for text in paragraphs:
        add_body(doc, text)

    add_subheading(doc, "1.2", "Identifikasi Masalah")
    for text in [
        "Informasi mengenai kekerasan, hak pengguna, dan pilihan bantuan belum selalu tersedia dalam format yang mudah dibaca dan kompatibel dengan berbagai teknologi asistif.",
        "Catatan pengalaman dan bukti sering tersimpan terpisah sehingga sulit disusun menjadi kronologi yang konsisten serta dapat diperiksa sumbernya.",
        "Proses formal dapat meminta pengguna menceritakan ulang informasi sensitif kepada beberapa pihak, meningkatkan beban emosional dan risiko hilangnya detail.",
        "Pengguna belum selalu memperoleh kendali granular mengenai data apa yang dipilih, diproses AI, dimasukkan ke laporan, atau dibagikan.",
        "AI generatif dapat menghasilkan teks yang tampak meyakinkan meskipun mengandung asumsi, sumber tidak valid, atau kesimpulan yang tidak diminta.",
    ]:
        add_bullet(doc, text)

    add_subheading(doc, "1.3", "Rumusan Masalah")
    questions = [
        "Bagaimana merancang alur edukasi, pencatatan, penataan bukti, dan persiapan laporan yang aksesibel serta tidak memaksa pengguna?",
        "Bagaimana menggunakan AI untuk mengurangi beban penyusunan kronologi tanpa menjadikannya penentu kebenaran, niat, diagnosis, atau status hukum?",
        "Bagaimana memastikan setiap hasil AI dapat ditelusuri ke catatan sumber, diedit, diterima, atau ditolak oleh pengguna?",
        "Bagaimana menyediakan dukungan percakapan yang empatik tetapi tetap non-klinis dan tidak menggantikan bantuan manusia?",
        "Bagaimana memvalidasi prototype secara jujur menggunakan data sintetis dan membedakan kemampuan saat ini dari visi sistem produksi?",
    ]
    for index, text in enumerate(questions, 1):
        add_bullet(doc, text, number=index)

    add_subheading(doc, "1.4", "Tujuan Proyek")
    goals = [
        "Membangun prototype web responsif yang menunjukkan alur AmanAkses dari edukasi hingga pratinjau laporan.",
        "Mengimplementasikan Safe Timeline Assistant dengan structured output, sumber wajib, fallback deterministik, dan human-review gate.",
        "Mengimplementasikan Asisten Aman sebagai chatbot navigasi dan dukungan emosional non-klinis dengan tool allowlist serta konfirmasi eksplisit.",
        "Menerapkan prinsip accessibility-by-design, privacy-by-design, consent-by-design, dan human-in-the-loop.",
        "Menguji fungsi inti melalui fixture sintetis, build, lint, audit rahasia, fallback, dan pemeriksaan aksesibilitas dasar.",
    ]
    for index, text in enumerate(goals, 1):
        add_bullet(doc, text, number=index)

    add_subheading(doc, "1.5", "Manfaat Proyek")
    add_minor_heading(doc, "Bagi Pengguna")
    add_body(
        doc,
        "AmanAkses memberi ruang untuk belajar dan menyusun informasi secara bertahap. Pengguna dapat memilih jalur yang sesuai kesiapan, meninjau kembali sumber, melewati AI, serta menghentikan proses kapan saja. Nilai utamanya bukan mendorong pengguna segera melapor, melainkan membantu pengguna memahami pilihan dan memegang kendali.",
    )
    add_minor_heading(doc, "Bagi Pendamping dan Lembaga")
    add_body(
        doc,
        "Rancangan laporan awal dapat membantu pendamping memahami urutan kejadian, sumber, kebutuhan akses, dan persetujuan pengguna. Meskipun integrasi eksternal belum tersedia, prototype menunjukkan bagaimana informasi dapat dipersiapkan secara lebih terstruktur tanpa menggantikan verifikasi manusia.",
    )
    add_minor_heading(doc, "Bagi Tim Pengembang")
    add_body(
        doc,
        "Proyek menjadi studi kasus penerapan AI yang tidak hanya mengejar kemampuan generatif, tetapi juga memperhatikan batas data, validasi skema, aksesibilitas, keselamatan emosional, kegagalan layanan, dan tanggung jawab manusia.",
    )

    add_subheading(doc, "1.6", "Target Pengguna")
    add_body(
        doc,
        "Target utama adalah mahasiswa atau masyarakat penyandang disabilitas yang memerlukan informasi dan ruang dokumentasi bertahap. Target sekunder meliputi pendamping tepercaya, Satgas PPKPT, lembaga layanan, bantuan hukum, dan pengelola layanan yang hanya dapat menerima data setelah persetujuan pengguna. Prototype tidak ditujukan untuk penggunaan kasus nyata sebelum melalui co-design, audit keamanan, evaluasi etika, dan pengujian bersama pengguna.",
    )

    add_subheading(doc, "1.7", "Batasan Masalah")
    limits = [
        "Seluruh data contoh, fixture, screenshot, dan feedback pada tugas ini bersifat sintetis.",
        "Prototype tidak menyediakan penyimpanan produksi, autentikasi nyata, enkripsi end-to-end, pengiriman laporan, atau integrasi lembaga eksternal.",
        "AI tidak menentukan kebenaran, niat, pelaku, kesalahan, diagnosis, tingkat risiko klinis, maupun keputusan hukum.",
        "Chatbot hanya memberi dukungan emosional non-klinis, navigasi, dan persiapan; bukan pengganti psikolog, psikiater, konselor, pendamping, atau layanan darurat.",
        "Validasi saat ini membuktikan fungsi teknis tertentu, bukan keamanan dan efektivitas penggunaan pada kasus nyata.",
    ]
    for text in limits:
        add_bullet(doc, text)


def add_chapter_two(doc):
    add_chapter(doc, "II", "ANALISIS MASALAH DAN STUDI LITERATUR")
    add_subheading(doc, "2.1", "Analisis Perjalanan Pengguna")
    add_body(
        doc,
        "Perjalanan pengguna tidak selalu dimulai dari keputusan untuk melapor. Pengguna dapat berada pada tahap memahami apakah suatu situasi membuatnya tidak aman, mencari informasi, menulis sebagian ingatan, mengumpulkan bukti, meminta pendapat pendamping, atau menyiapkan percakapan dengan layanan. Karena itu, solusi yang langsung memaksa formulir final dapat menambah beban dan tidak sesuai dengan kesiapan pengguna.",
    )
    journey_rows = [
        ("Memahami", "Bahasa teknis, konten panjang, format tidak aksesibel", "Easy-read, audio/teks alternatif, navigasi jelas"),
        ("Mencatat", "Takut kehilangan draf dan harus menyelesaikan sekaligus", "Penyimpanan bertahap, kendali hapus, struktur opsional"),
        ("Menata bukti", "File tersebar dan konteks mudah hilang", "Label, keterkaitan sumber, status sintetis yang jelas"),
        ("Menyusun kronologi", "Urutan waktu sulit dan pengguna perlu mengulang cerita", "Draf bersumber, field kosong bila tidak diketahui"),
        ("Mencari bantuan", "Tidak tahu pihak yang sesuai atau data yang akan diterima", "Pilihan bantuan, kebutuhan akses, consent eksplisit"),
    ]
    add_table(
        doc,
        ["Tahap", "Hambatan Utama", "Kebutuhan Sistem"],
        journey_rows,
        widths=[3.2, 5.4, 6.0],
    )

    add_subheading(doc, "2.2", "Akar Masalah")
    causes = [
        "Aksesibilitas sering diperlakukan sebagai fitur tambahan, bukan persyaratan arsitektural sejak awal.",
        "Sistem pelaporan berorientasi pada kelengkapan administrasi, sementara pengguna dapat membutuhkan tahap persiapan dan jeda.",
        "Catatan, bukti, kebutuhan akses, dan persetujuan berada pada alat yang terpisah.",
        "Ketergantungan pada teks bebas membuat penyusunan ulang kronologi memerlukan beban ingatan dan kognitif yang tinggi.",
        "Output AI yang rapi dapat memicu automation bias apabila tidak disertai sumber, ketidakpastian, dan keputusan eksplisit.",
    ]
    for text in causes:
        add_bullet(doc, text)

    add_subheading(doc, "2.3", "Studi Literatur dan Landasan Konsep")
    add_minor_heading(doc, "Artificial Intelligence dan Large Language Model")
    add_body(
        doc,
        "Artificial Intelligence adalah bidang yang mengembangkan sistem untuk melakukan tugas yang biasanya membutuhkan kemampuan kognitif manusia. Large Language Model merupakan salah satu pendekatan AI generatif yang dapat memahami instruksi dan menghasilkan teks. Dalam AmanAkses, LLM tidak digunakan sebagai evaluator kasus. Model hanya mengubah input terpilih menjadi struktur terbatas yang dapat diperiksa.",
    )
    add_minor_heading(doc, "Human-in-the-Loop")
    add_body(
        doc,
        "Human-in-the-loop berarti manusia bukan sekadar penerima hasil, tetapi memiliki titik keputusan yang nyata. Pada Safe Timeline Assistant, setiap peristiwa dimulai sebagai draf, menampilkan sumber, dapat diedit, dan harus diterima atau ditolak. Pada Asisten Aman, tool tidak berjalan otomatis dan selalu menunggu konfirmasi.",
    )
    add_minor_heading(doc, "Structured Output dan Validasi")
    add_body(
        doc,
        "Structured output membatasi keluaran Gemini ke kontrak JSON yang telah ditentukan. Batas ini meningkatkan konsistensi, tetapi tidak menjamin kebenaran semantik. Karena itu, aplikasi memeriksa tipe data, nilai enum, keberadaan sourceNoteIds, route internal, nama tool, dan status review sebelum hasil ditampilkan.",
    )
    add_minor_heading(doc, "Aksesibilitas Digital")
    add_body(
        doc,
        "WCAG 2.2 merumuskan empat prinsip utama: perceivable, operable, understandable, dan robust. Prinsip tersebut diterjemahkan menjadi kontras yang memadai, label semantik, navigasi keyboard, fokus yang terlihat, ukuran target yang cukup, bahasa ringkas, dukungan pengurangan gerak, dan layout responsif.",
    )
    add_minor_heading(doc, "Trauma-Informed dan Consent-by-Design")
    add_body(
        doc,
        "Desain berbasis trauma menghindari bahasa menyalahkan, detail grafis yang tidak diperlukan, tekanan untuk menyelesaikan alur, dan tindakan yang terjadi tanpa persetujuan. Consent-by-design berarti pengguna mengetahui data apa yang dipilih, apa yang diproses, dan tindakan apa yang akan dilakukan sebelum memberikan persetujuan.",
    )

    add_subheading(doc, "2.4", "Persona dan Kebutuhan Akses")
    persona_rows = [
        ("Netra/low vision", "Pembaca layar, struktur heading, keyboard, kontras, teks alternatif"),
        ("Tuli/hambatan pendengaran", "Informasi visual dan tekstual, caption, notifikasi non-audio"),
        ("Kognitif/neurodivergent", "Bahasa sederhana, langkah kecil, konsistensi, ringkasan"),
        ("Fisik/motorik", "Target besar, navigasi keyboard, interaksi minim presisi"),
        ("Psikososial", "Pilihan jeda, safe exit, bahasa tenang, tidak memaksa"),
        ("Pendamping tepercaya", "Konteks terstruktur, kebutuhan akses, batas persetujuan"),
        ("Satgas/lembaga", "Laporan awal terstruktur dan status sumber yang dapat diperiksa"),
    ]
    add_table(doc, ["Persona", "Kebutuhan Utama"], persona_rows, widths=[4.2, 10.4])
    add_caption(doc, "Tabel 2.1 Persona dan kebutuhan akses")

    add_subheading(doc, "2.5", "Kebutuhan Fungsional")
    functional_rows = [
        ("F-01", "Pahami Kekerasan", "Materi easy-read, checklist, dan jalur menuju bantuan"),
        ("F-02", "Jurnal Aman", "Membuat dan mengubah catatan secara bertahap"),
        ("F-03", "Brankas Bukti", "Menata file dan metadata sintetis"),
        ("F-04", "Safe Timeline Assistant", "Menyusun draf kronologi dari catatan terpilih"),
        ("F-05", "Review Peristiwa", "Membuka sumber, mengedit, menerima, atau menolak"),
        ("F-06", "Asisten Aman", "Dukungan non-klinis, navigasi, dan persiapan"),
        ("F-07", "Tool Terkonfirmasi", "Timeline, jurnal, aksesibilitas, dan bantuan manusia"),
        ("F-08", "Laporan Awal", "Pratinjau informasi yang dipilih pengguna"),
        ("F-09", "Aksesibilitas", "Easy-read, ukuran teks, kontras, pengurangan gerak"),
        ("F-10", "Keluar Cepat", "Meninggalkan alur sensitif menuju halaman netral"),
    ]
    add_table(doc, ["Kode", "Fitur", "Deskripsi"], functional_rows, widths=[1.5, 4.3, 8.8])
    add_caption(doc, "Tabel 2.2 Kebutuhan fungsional prioritas")

    add_subheading(doc, "2.6", "Kebutuhan Non-Fungsional")
    nonfunctional_rows = [
        ("Aksesibilitas", "Mengikuti prinsip WCAG 2.2 dan dapat dioperasikan dengan keyboard"),
        ("Privasi", "Minimisasi data, key server-side, data sintetis untuk demo"),
        ("Keandalan", "Fallback deterministik ketika layanan model gagal"),
        ("Auditabilitas", "Sumber wajib, status ketidakpastian, dan status review"),
        ("Keamanan emosional", "Bahasa netral, pilihan jeda, tidak meminta detail berlebihan"),
        ("Transparansi", "Label draf, mode Gemini/fallback, dan batas prototype"),
        ("Maintainability", "TypeScript, schema bersama, fixture, lint, dan build"),
    ]
    add_table(doc, ["Aspek", "Kebutuhan"], nonfunctional_rows, widths=[4.0, 10.6])
    add_caption(doc, "Tabel 2.3 Kebutuhan non-fungsional")

    add_subheading(doc, "2.7", "Analisis Sistem Sejenis dan Kesenjangan")
    add_body(
        doc,
        "Catatan ponsel dapat menyimpan teks, tetapi tidak menggabungkan kebutuhan akses, sumber, consent, dan laporan. Folder media menyimpan screenshot, tetapi konteks dan urutan mudah hilang. Form aduan formal menyediakan kanal, tetapi dapat mengasumsikan pengguna sudah siap mengisi semua informasi. Chatbot umum mudah digunakan, tetapi dapat memberikan jawaban terlalu luas dan tidak selalu memiliki guardrail untuk konteks sensitif. AmanAkses mencoba mengisi kesenjangan tersebut melalui alur terpadu, AI yang dibatasi, serta keputusan eksplisit pada pengguna.",
    )


def add_chapter_three(doc):
    add_chapter(doc, "III", "PEMANFAATAN AI DALAM PENGEMBANGAN SOLUSI")
    add_subheading(doc, "3.1", "Klasifikasi Penggunaan AI")
    add_body(
        doc,
        "Penggunaan AI dalam proyek dibedakan menjadi dua kelompok. Pertama, AI di dalam produk, yaitu Gemini yang mendukung Safe Timeline Assistant dan Asisten Aman. Kedua, AI sebagai partner kerja tim, yaitu ChatGPT dan Codex yang membantu analisis, perancangan, implementasi, debugging, pengujian, serta dokumentasi. Alat seperti Microsoft Word, python-docx, PowerPoint, Vite, TypeScript, dan ESLint merupakan alat produksi non-AI.",
    )
    tools_rows = [
        ("Gemini", "AI di dalam produk", "Timeline terstruktur dan respons chatbot", "Halusinasi, inferensi, route/tool salah", "Schema, allowlist, sumber, review"),
        ("ChatGPT", "Partner tim", "Brainstorming dan kritik konsep", "Jawaban generik atau klaim berlebih", "Dibandingkan dengan panduan dan sumber"),
        ("Codex", "Partner tim", "Inspeksi kode, implementasi, QA, dokumen", "Perubahan tidak sesuai codebase", "Diff, build, lint, fixture, render"),
        ("Word/python-docx", "Non-AI", "Penyusunan dan ekspor dokumen", "Masalah layout", "Render tiap halaman dan inspeksi"),
        ("Vite/ESLint", "Non-AI", "Runtime lokal dan pemeriksaan statis", "Tidak menguji makna produk", "Dikombinasikan dengan fixture dan QA"),
    ]
    add_table(
        doc,
        ["Alat", "Kategori", "Keluaran", "Risiko", "Validasi Manusia"],
        tools_rows,
        widths=[2.2, 2.6, 3.5, 3.2, 3.4],
        font_size=9,
    )
    add_caption(doc, "Tabel 3.1 Pemetaan AI dan alat produksi")

    add_subheading(doc, "3.2", "Proses Brainstorming Menggunakan AI")
    add_body(
        doc,
        "Brainstorming dimulai dari pembacaan panduan tugas dan inventarisasi masalah, bukan dari pemilihan model. AI diminta membandingkan ide AmanAkses dengan kriteria tugas, menjelaskan letak dampak, dan mengkritik fitur yang terlalu luas. Hasil awal menyarankan banyak kemampuan generatif, tetapi tim mempersempit ruang lingkup menjadi timeline bersumber dan chatbot non-klinis. Keputusan tersebut mengurangi risiko sekaligus membuat kontribusi AI dapat didemonstrasikan dengan jelas.",
    )
    collaboration_steps = [
        "Tim memberi konteks berupa panduan tugas, tujuan produk, pengguna sasaran, codebase, dan batas etika.",
        "AI menghasilkan alternatif struktur, risiko, atau diagnosis teknis.",
        "Tim memeriksa hasil terhadap sumber resmi, kemampuan prototype, dan dampak pada pengguna.",
        "Perubahan diterapkan secara terbatas, kemudian diuji melalui build, lint, fixture, browser QA, atau render.",
        "Kegagalan dicatat dan dipakai untuk memperbaiki prompt, validator, fallback, atau antarmuka.",
    ]
    for index, text in enumerate(collaboration_steps, 1):
        add_bullet(doc, text, number=index)

    add_subheading(doc, "3.3", "Dokumentasi Prompt")
    prompt_rows = [
        ("Analisis masalah", "Bandingkan AmanAkses dengan panduan dan cari masalah nyata yang paling kuat.", "AI terlalu cepat memperluas fitur.", "Fokus pada alur dokumentasi dan aksesibilitas."),
        ("Desain AI", "Tentukan fungsi AI yang relevan dan batas yang tidak boleh dilanggar.", "Timeline dipilih; klasifikasi kasus ditolak.", "AI hanya menyusun draf."),
        ("Implementasi", "Tambahkan endpoint Gemini server-side dengan structured JSON dan fallback.", "Struktur awal belum memaksa sumber.", "sourceNoteIds diwajibkan."),
        ("Debugging", "Cari penyebab mode fallback meskipun .env tersedia.", "Konfigurasi dibaca pada proses yang salah.", "API lokal dijalankan bersama Vite."),
        ("Chatbot", "Buat chatbot suportif yang dapat memakai tool internal secara aman.", "Risiko dianggap sebagai konselor.", "Non-klinis, allowlist, konfirmasi."),
        ("Validasi", "Buat fixture tanggal ambigu, field hilang, source palsu, dan API gagal.", "Kasus negatif belum terpisah.", "10 skenario + 1 negative guard."),
        ("Dokumentasi", "Tulis laporan akademik tanpa mengklaim fitur produksi atau peer test nyata.", "Klaim teknis lama tidak akurat.", "Status prototype dan visi dipisahkan."),
    ]
    add_table(
        doc,
        ["Tahap", "Inti Prompt", "Masalah Hasil Awal", "Keputusan/Revisi"],
        prompt_rows,
        widths=[2.2, 5.0, 3.7, 4.0],
        font_size=9,
    )
    add_caption(doc, "Tabel 3.2 Log prompt dan keputusan manusia")

    add_subheading(doc, "3.4", "AI di Dalam Produk: Safe Timeline Assistant")
    add_body(
        doc,
        "Safe Timeline Assistant menerima array catatan yang dipilih pengguna. Endpoint membaca GEMINI_API_KEY pada sisi server dan meminta keluaran JSON terstruktur. Setiap event berisi id, tanggal, waktu, lokasi, ringkasan, sourceNoteIds, ketidakpastian, dan requiresReview. Field yang tidak disebutkan harus bernilai null. Semua event wajib berstatus draf sampai pengguna mengambil keputusan.",
    )
    add_minor_heading(doc, "System Instruction Inti")
    for text in [
        "Gunakan hanya fakta yang tertulis eksplisit pada catatan terpilih.",
        "Jangan menyimpulkan niat, kesalahan, kebenaran, diagnosis, atau status hukum.",
        "Jangan mengisi tanggal, waktu, lokasi, atau identitas yang tidak tersedia.",
        "Gunakan null untuk informasi yang hilang dan tandai ketidakpastian.",
        "Setiap peristiwa wajib memiliki sourceNoteIds yang valid.",
        "Gunakan bahasa Indonesia yang netral, ringkas, dan tidak grafis.",
        "requiresReview selalu true.",
    ]:
        add_bullet(doc, text)

    add_minor_heading(doc, "Kontrak Structured Output")
    schema_rows = [
        ("id", "string", "ID event unik"),
        ("date", "string atau null", "Tanggal hanya jika eksplisit"),
        ("time", "string atau null", "Waktu hanya jika eksplisit"),
        ("location", "string atau null", "Lokasi hanya jika eksplisit"),
        ("summary", "string", "Ringkasan netral dari sumber"),
        ("sourceNoteIds", "array string", "Wajib merujuk catatan input"),
        ("uncertainty", "explicit/ambiguous/missing", "Status kepastian informasi"),
        ("requiresReview", "boolean", "Selalu true"),
    ]
    add_table(doc, ["Field", "Tipe", "Aturan"], schema_rows, widths=[3.3, 4.0, 7.3])

    add_subheading(doc, "3.5", "AI di Dalam Produk: Asisten Aman")
    add_body(
        doc,
        "Asisten Aman adalah chatbot untuk dukungan emosional non-klinis, navigasi, dan persiapan. Ia dapat menjelaskan menu, membantu memilih langkah kecil, menyusun kerangka jurnal, membuka pengaturan aksesibilitas, atau membuat handoff ke timeline. Chatbot tidak meminta pengguna menceritakan detail kasus, tidak memberi diagnosis atau terapi, dan tidak menghubungi pihak luar.",
    )
    for text in [
        "draft_timeline: menyiapkan catatan sintetis dan membuka review Safe Timeline Assistant.",
        "prepare_journal: mengisi editor dengan kerangka yang dapat diubah atau dihapus.",
        "update_accessibility: menerapkan Easy Read, ukuran teks, dan pengurangan gerak.",
        "open_support: membuka pendamping, pusat bantuan, atau safe exit tanpa mengirim data.",
    ]:
        add_bullet(doc, text)
    add_body(
        doc,
        "Semua tool menggunakan allowlist dan memerlukan konfirmasi eksplisit. Tool yang tidak dikenal dibuang. Dalam respons darurat, tool produktivitas seperti timeline dan jurnal diblokir agar antarmuka memprioritaskan bantuan manusia dan keselamatan.",
    )

    evaluation_heading = add_subheading(
        doc,
        "3.6",
        "Evaluasi Hasil AI, Kesalahan, dan Koreksi",
    )
    evaluation_heading.paragraph_format.page_break_before = True
    error_rows = [
        ("Tanggal/waktu diasumsikan", "Model memakai waktu metadata sebagai waktu kejadian", "Field menjadi null jika tidak tertulis"),
        ("Sumber tidak valid", "Event merujuk ID yang tidak ada", "Seluruh respons ditolak validator"),
        ("Kesimpulan substantif", "Ringkasan menyiratkan suatu pelanggaran pasti terjadi", "Prompt netral dan larangan kesimpulan"),
        ("Ketidakpastian keliru", "Frasa seperti “seingatku” dianggap eksplisit", "Fixture ambigu dan revisi aturan"),
        ("Route eksternal", "Chatbot mengusulkan URL di luar aplikasi", "Route allowlist internal"),
        ("Tool saat darurat", "Chatbot menawarkan jurnal/timeline", "Guard memblokir tool produktivitas"),
    ]
    add_caption(doc, "Tabel 3.3 Kesalahan AI dan koreksi")
    add_table(
        doc,
        ["Kesalahan", "Risiko", "Koreksi"],
        error_rows,
        widths=[3.4, 5.2, 6.0],
        font_size=9,
    )

    add_image(
        doc,
        DIAGRAMS / "AmanAkses_Arsitektur_Sistem.png",
        "Gambar 3.1 Arsitektur sistem AmanAkses",
        width_cm=14.7,
    )

    add_subheading(doc, "3.7", "Prinsip Validasi dan Tanggung Jawab")
    add_body(
        doc,
        "Prompt bukan mekanisme keamanan tunggal. AmanAkses menggabungkan pembatasan input, structured output, validator, fallback konservatif, label draf, sumber, kontrol edit/terima/tolak, konfirmasi tool, dan pengujian. Meskipun demikian, tanggung jawab tetap berada pada tim dan pengguna manusia. Hasil AI tidak boleh digunakan sebagai bukti kebenaran atau dasar tindakan eksternal otomatis.",
    )


def add_chapter_four(doc):
    add_chapter(doc, "IV", "PERANCANGAN SOLUSI")
    add_subheading(doc, "4.1", "Nama dan Konsep Produk")
    add_body(
        doc,
        "AmanAkses adalah platform dokumentasi aman dan aksesibel yang dirancang sebagai ruang persiapan. Nilai produk dirangkum dalam empat prinsip: Understand Safely, Record Safely, Decide Safely, dan Share Safely. Pengguna dapat memulai dari edukasi, mencatat saat siap, menata informasi, meminta bantuan AI secara opsional, meninjau hasil, dan menentukan langkah berikutnya.",
    )

    add_subheading(doc, "4.2", "Keunikan dan Nilai Inovasi")
    for text in [
        "Menggabungkan aksesibilitas multi-disabilitas dengan dokumentasi, AI bersumber, dan consent dalam satu perjalanan.",
        "Tidak menganggap pelaporan sebagai satu-satunya tujuan; belajar, mencatat, dan meminta pendampingan juga merupakan hasil yang valid.",
        "Menempatkan AI pada tugas sempit yang dapat diaudit, bukan pada penilaian kasus.",
        "Menyediakan jalur fallback dan jalur manual sehingga fungsi utama tidak sepenuhnya bergantung pada model.",
        "Memisahkan kemampuan prototype dari visi produksi secara transparan.",
    ]:
        add_bullet(doc, text)

    add_subheading(doc, "4.3", "Fitur Utama dan Status Prototype")
    feature_rows = [
        ("Dashboard", "Memilih langkah sesuai kesiapan", "Tersedia dengan data sintetis"),
        ("Pahami Kekerasan", "Edukasi easy-read dan checklist", "Tersedia"),
        ("Jurnal Aman", "Catatan bertahap dan kerangka", "Tersedia sebagai simulasi UI"),
        ("Brankas Bukti", "Menata bukti dan metadata", "Tersedia sebagai data sintetis"),
        ("Safe Timeline", "Draf kronologi bersumber", "Tersedia; Gemini/fallback"),
        ("Asisten Aman", "Navigasi dan dukungan non-klinis", "Tersedia; tool terkonfirmasi"),
        ("Pendamping", "Pilihan bantuan manusia", "Simulasi UI"),
        ("Laporan Awal", "Pratinjau data terpilih", "Simulasi UI"),
        ("Aksesibilitas", "Easy-read, ukuran teks, reduce motion", "Tersedia pada state sesi"),
        ("Keluar Cepat", "Meninggalkan halaman sensitif", "Tersedia"),
    ]
    add_table(
        doc,
        ["Fitur", "Manfaat", "Status"],
        feature_rows,
        widths=[3.3, 6.5, 4.8],
        font_size=9,
    )
    add_caption(doc, "Tabel 4.1 Fitur, manfaat, dan status prototype")

    add_subheading(doc, "4.4", "Workflow System")
    workflow = [
        "Pengguna membuka dashboard dan memilih titik awal: belajar, mencatat, menyusun kronologi, atau mencari bantuan.",
        "Pengguna mengaktifkan preferensi aksesibilitas sesuai kebutuhan.",
        "Pengguna membaca materi easy-read atau membuat jurnal menggunakan data sintetis.",
        "Pengguna memilih catatan yang diizinkan untuk diproses oleh Safe Timeline Assistant.",
        "Gemini atau fallback menghasilkan draf terstruktur yang divalidasi.",
        "Pengguna membuka sumber, mengedit, menerima, atau menolak setiap event.",
        "Asisten Aman dapat membantu navigasi atau menawarkan tool yang harus dikonfirmasi.",
        "Pengguna meninjau laporan awal dan consent; prototype tidak mengirim data keluar.",
        "Pengguna dapat menekan Keluar Cepat kapan saja.",
    ]
    for index, text in enumerate(workflow, 1):
        add_bullet(doc, text, number=index)

    add_image(
        doc,
        DIAGRAMS / "AmanAkses_Alur_Pengguna_AI.png",
        "Gambar 4.1 Alur pengguna dan kendali AI",
        width_cm=14.7,
    )

    add_subheading(doc, "4.5", "Desain Antarmuka dan Aksesibilitas")
    add_body(
        doc,
        "Antarmuka menggunakan warna teal, violet, warm-white, dan slate dengan hierarki visual yang konsisten. Teks utama memakai ukuran yang cukup, kartu memiliki ruang yang tidak padat, dan tombol aksi utama dibedakan secara jelas. Struktur halaman memanfaatkan heading semantik, label tombol, indikator fokus, dan komponen yang dapat dinavigasi dengan keyboard.",
    )
    for text in [
        "Perceivable: kontras, teks alternatif, status non-warna, dan konten easy-read.",
        "Operable: target interaksi yang cukup besar, keyboard, fokus terlihat, dan reduce motion.",
        "Understandable: bahasa singkat, tindakan eksplisit, langkah bertahap, dan label draf.",
        "Robust: komponen React semantik dan validasi agar data tidak merusak antarmuka.",
    ]:
        add_bullet(doc, text)

    add_subheading(doc, "4.6", "Arsitektur Teknis dan Data Boundary")
    add_body(
        doc,
        "Prototype menggunakan React, TypeScript, Vite, React Router, Tailwind CSS, dan Motion pada antarmuka. Vite local API middleware menjalankan endpoint /api/timeline dan /api/chat dalam proses pengembangan yang sama. Pada deployment yang mendukung fungsi Node server-side, endpoint dapat dijalankan sebagai serverless function. GEMINI_API_KEY hanya dibaca pada server dan tidak menggunakan awalan VITE_.",
    )
    tech_rows = [
        ("Frontend", "React, TypeScript, Vite, Tailwind", "Terimplementasi"),
        ("Routing", "React Router", "Terimplementasi"),
        ("AI API", "Gemini structured JSON", "Terimplementasi, membutuhkan key"),
        ("Fallback", "Ekstraksi deterministik lokal", "Terimplementasi"),
        ("Validasi", "Shared schema dan allowlist", "Terimplementasi"),
        ("Penyimpanan", "State/data sintetis", "Belum produksi"),
        ("Autentikasi", "Simulasi pengalaman", "Belum produksi"),
        ("Enkripsi", "Rancangan kebutuhan", "Belum diimplementasikan"),
        ("Integrasi lembaga", "Rancangan consent dan laporan", "Belum diimplementasikan"),
    ]
    add_table(
        doc,
        ["Lapisan", "Teknologi/Pendekatan", "Status"],
        tech_rows,
        widths=[3.2, 6.8, 4.6],
        font_size=9,
    )
    add_caption(doc, "Tabel 4.2 Teknologi prototype dan visi produksi")

    add_subheading(doc, "4.7", "Visi Arsitektur Produksi")
    add_body(
        doc,
        "Sistem produksi memerlukan autentikasi yang aman, penyimpanan terenkripsi, object storage, consent ledger, audit log, kebijakan retensi, penghapusan data, monitoring, incident response, dan pengujian keamanan independen. Integrasi dengan lembaga harus didahului verifikasi SOP, perjanjian pengelolaan data, serta co-design dengan organisasi disabilitas dan penyedia layanan. Komponen tersebut merupakan roadmap, bukan kemampuan prototype saat ini.",
    )


def add_chapter_five(doc):
    add_chapter(doc, "V", "IMPLEMENTASI PROTOTYPE")
    add_subheading(doc, "5.1", "Bentuk dan Lingkungan Pengembangan")
    add_body(
        doc,
        "Prototype berbentuk aplikasi web responsif yang dapat dijalankan pada desktop dan perangkat mobile. Aplikasi dikembangkan dengan Node.js dan pnpm. Perintah pnpm dev menjalankan frontend serta API lokal, sehingga Vercel CLI tidak diperlukan untuk demonstrasi lokal. Deployment publik bersifat opsional.",
    )
    environment = [
        "Sistem operasi pengembangan: Windows.",
        "Runtime: Node.js dan pnpm.",
        "Frontend: React, TypeScript, Vite, Tailwind CSS, React Router, Motion.",
        "API lokal: middleware Vite untuk /api/timeline dan /api/chat.",
        "AI: Gemini melalui GEMINI_API_KEY di server; model dapat diatur melalui GEMINI_MODEL.",
        "Quality assurance: TypeScript build, ESLint, fixture timeline, fixture chatbot, dan browser QA.",
    ]
    for text in environment:
        add_bullet(doc, text)

    add_subheading(doc, "5.2", "Implementasi Pahami Kekerasan")
    add_body(
        doc,
        "Halaman Pahami Kekerasan berisi modul edukasi non-grafis, tujuan belajar, mode ringkas dan lengkap, checklist, kalimat bantu, progres, serta jalur menuju jurnal atau bantuan manusia. Konten tidak memaksa pengguna menentukan label terhadap pengalaman. Fokusnya adalah membantu pengguna mengenali batas pribadi, pilihan aman, dan sumber dukungan.",
    )
    add_image(
        doc,
        SCREENSHOTS / "pahami-kekerasan.png",
        "Gambar 5.1 Halaman Pahami Kekerasan",
        width_cm=14.7,
    )

    add_subheading(doc, "5.3", "Implementasi Asisten Aman")
    add_body(
        doc,
        "Asisten Aman menggunakan respons terstruktur yang memisahkan pesan, tingkat keselamatan, route internal, saran tindak lanjut, dan usulan tool. Pada percakapan biasa, asisten dapat membantu pengguna mengecilkan pilihan menjadi satu langkah. Pada bahasa yang menunjukkan keadaan darurat, asisten memprioritaskan bantuan manusia dan memblokir tool produktivitas.",
    )
    add_image(
        doc,
        SCREENSHOTS / "chatbot-tools.png",
        "Gambar 5.2 Asisten Aman dan tool terstruktur",
        width_cm=14.7,
    )

    add_subheading(doc, "5.4", "Implementasi Safe Timeline Assistant")
    add_body(
        doc,
        "Halaman kronologi menyediakan skenario sintetis, daftar catatan sumber, metadata waktu, bukti terkait, kebutuhan akses, serta workspace untuk memilih input. Setelah tombol penyusunan dijalankan, aplikasi menampilkan mode Gemini atau fallback dan membuat draf event. Pengguna dapat berpindah dari workspace ke review tanpa kehilangan konteks.",
    )
    add_image(
        doc,
        SCREENSHOTS / "kronologi-workspace.png",
        "Gambar 5.3 Workspace Safe Timeline Assistant",
        width_cm=14.7,
    )
    add_body(
        doc,
        "Pada tahap review, setiap event menampilkan sumber yang dapat dibuka, field yang tidak pasti, editor, dan tiga status keputusan: pending, accepted, atau rejected. Hanya event accepted yang dinyatakan siap dipertimbangkan dalam laporan. Perubahan status tidak dikirim ke pihak luar.",
    )
    add_image(
        doc,
        SCREENSHOTS / "kronologi-review.png",
        "Gambar 5.4 Review peristiwa dan referensi sumber",
        width_cm=14.7,
    )

    add_subheading(doc, "5.5", "Fallback dan Penanganan Kegagalan API")
    add_body(
        doc,
        "Fallback deterministik memastikan demonstrasi tetap berjalan ketika API key tidak tersedia, jaringan gagal, kuota habis, respons model tidak valid, atau endpoint mengalami kesalahan. Fallback hanya membaca label eksplisit seperti Tanggal, Waktu, dan Lokasi dari data sintetis. Pendekatan ini bukan pengganti NLP penuh, tetapi sengaja konservatif agar tidak menebak.",
    )

    add_subheading(doc, "5.6", "Konfigurasi Rahasia")
    add_body(
        doc,
        "GEMINI_API_KEY ditempatkan pada .env.local atau environment server dan tidak boleh diberi awalan VITE_. Variabel dengan awalan VITE_ dapat masuk ke bundle browser. SourceCode.zip dan repository pengumpulan harus mengecualikan .env, .env.local, log, node_modules, dan build output. .env.example hanya memuat nama variabel tanpa nilai rahasia.",
    )

    add_subheading(doc, "5.7", "Skenario Demonstrasi Manual")
    demo = [
        "Buka dashboard dan jelaskan bahwa pengguna dapat memilih langkah tanpa dipaksa melapor.",
        "Masuk ke Pahami Kekerasan dan tunjukkan mode ringkas, checklist, serta navigasi menuju bantuan.",
        "Buka Asisten Aman, kirim pertanyaan navigasi, lalu konfirmasi satu tool aksesibilitas atau jurnal.",
        "Masuk ke Safe Timeline Assistant dan pilih tiga catatan sintetis.",
        "Jalankan penyusunan, tunjukkan mode Gemini/fallback, serta jelaskan bahwa hasil adalah draf.",
        "Buka source note, edit satu field, terima satu event, dan tolak satu event.",
        "Tunjukkan laporan awal, consent, pengaturan aksesibilitas, dan Keluar Cepat.",
        "Akhiri dengan batas sistem: tidak ada diagnosis, kebenaran otomatis, atau pengiriman eksternal.",
    ]
    for index, text in enumerate(demo, 1):
        add_bullet(doc, text, number=index)

    add_subheading(doc, "5.8", "Keterbatasan Implementasi")
    limitations = [
        "Penyimpanan, autentikasi, enkripsi, upload bukti, berbagi, dan ekspor masih simulasi UI.",
        "Fallback hanya mengenali pola eksplisit pada data contoh.",
        "Gemini tetap dapat salah memahami konteks meskipun schema dibatasi.",
        "Deteksi bahasa darurat pada chatbot bukan penilaian risiko klinis.",
        "Direktori bantuan dan identitas layanan merupakan contoh.",
        "Prototype belum diuji bersama penyandang disabilitas atau penyintas.",
    ]
    for text in limitations:
        add_bullet(doc, text)


def add_chapter_six(doc):
    add_chapter(doc, "VI", "VALIDASI DAN EVALUASI")
    add_subheading(doc, "6.1", "Metode Pengujian")
    add_body(
        doc,
        "Validasi dilakukan berlapis. Pemeriksaan statis memastikan kode dapat dibangun dan mengikuti aturan lint. Fixture timeline memeriksa struktur, sumber, field hilang, ketidakpastian, dan fallback. Fixture chatbot memeriksa navigasi, dukungan emosional, route, tool, konfirmasi, serta kondisi darurat. Browser QA memeriksa tampilan desktop dan mobile, keyboard dasar, overflow, dan console. Audit rahasia memastikan API key tidak masuk bundle atau arsip.",
    )

    add_subheading(doc, "6.2", "Hasil Pengujian Dasar")
    validation_rows = [
        ("Build produksi", "pnpm build", "Lulus"),
        ("Lint", "pnpm lint", "Lulus"),
        ("Timeline utama", "10 skenario sintetis", "10/10 lulus"),
        ("Negative source guard", "sourceNoteIds tidak valid", "Lulus; respons ditolak"),
        ("Chatbot", "12 skenario navigasi/tool/keselamatan", "12/12 lulus"),
        ("Route eksternal", "Route di luar allowlist", "Ditolak"),
        ("Tool tidak dikenal", "Nama tool di luar allowlist", "Dibuang"),
        ("Darurat", "Tool produktivitas saat urgent", "Diblokir"),
        ("Fallback", "API/key tidak tersedia", "Berfungsi"),
        ("Secret audit", ".env dan browser bundle", "Tidak ditemukan key"),
        ("Responsive QA", "Desktop dan viewport 390 px", "Tanpa overflow/error console"),
        ("Keyboard dasar", "Navigasi halaman utama", "Diperiksa"),
    ]
    add_table(
        doc,
        ["Pemeriksaan", "Bukti/Metode", "Hasil"],
        validation_rows,
        widths=[4.0, 6.0, 4.6],
        font_size=9,
    )
    add_caption(doc, "Tabel 6.1 Hasil validasi teknis")
    add_body(
        doc,
        "Hasil timeline dirumuskan sebagai “10 skenario utama lulus + 1 negative guard lulus”. Pemisahan ini penting karena negative guard tidak menguji kemampuan menghasilkan event, melainkan kemampuan sistem menolak respons yang berbahaya atau tidak bersumber.",
    )

    add_subheading(doc, "6.3", "Skenario Validasi Timeline")
    scenarios = [
        "Tanggal, waktu, dan lokasi eksplisit dipertahankan.",
        "Waktu yang tidak tersedia tetap null.",
        "Lokasi yang tidak tersedia tetap null.",
        "Tanggal relatif atau ambigu ditandai uncertainty.",
        "Beberapa catatan dapat mendukung satu event.",
        "Setiap event hanya merujuk sourceNoteIds input.",
        "Urutan event tidak mengubah fakta sumber.",
        "Ringkasan tidak menambahkan niat atau kesimpulan hukum.",
        "requiresReview selalu true.",
        "Fallback menghasilkan struktur yang dapat ditinjau.",
        "Negative guard menolak sumber yang tidak valid.",
    ]
    for index, text in enumerate(scenarios, 1):
        add_bullet(doc, text, number=index)

    add_subheading(doc, "6.4", "Feedback Pengguna Sintetis")
    add_body(
        doc,
        "Workbook Data_Feedback_Pengguna_Sintetis.xlsx berisi 30 respons buatan untuk menguji struktur instrumen, rumus, dan visualisasi. Data ini bukan hasil survei, bukan peer testing, dan tidak dapat digunakan untuk membuktikan kepuasan atau dampak. Nilainya hanya membantu tim mempersiapkan pertanyaan dan hipotesis sebelum pengujian nyata.",
    )
    feedback_rows = [
        ("Jumlah respons", "30 respons sintetis"),
        ("Keberhasilan skenario", "27 berhasil; 3 perlu bantuan"),
        ("Waktu rata-rata", "5,49 menit"),
        ("Kemudahan", "4,20 dari 5"),
        ("Kejelasan", "4,53 dari 5"),
        ("Aksesibilitas", "4,50 dari 5"),
        ("Kepercayaan terhadap AI", "4,00 dari 5"),
        ("Kendali manusia", "4,50 dari 5"),
        ("Kegunaan timeline", "4,50 dari 5"),
        ("Kegunaan chatbot", "4,00 dari 5"),
        ("Minat menggunakan", "26 Ya; 4 Mungkin"),
    ]
    add_table(doc, ["Indikator", "Nilai Simulasi"], feedback_rows, widths=[7.0, 7.6])
    add_caption(doc, "Tabel 6.2 Ringkasan data feedback sintetis")
    add_image(
        doc,
        FEEDBACK_PREVIEW / "dashboard.png",
        "Gambar 6.1 Dashboard feedback pengguna sintetis",
        width_cm=14.7,
    )
    add_body(
        doc,
        "Hipotesis yang muncul dari data simulasi adalah perlunya penjelasan fallback yang lebih singkat, label yang semakin jelas bahwa chatbot bukan konselor, serta penyederhanaan beberapa kartu. Hipotesis tersebut wajib diuji kembali dengan 3–5 mahasiswa menggunakan skenario sintetis sebelum dimasukkan sebagai temuan nyata.",
    )

    add_subheading(doc, "6.5", "Rencana Peer Testing")
    add_body(
        doc,
        "Peer testing dilakukan tanpa meminta pengalaman pribadi peserta. Setiap peserta menggunakan data sintetis, menyelesaikan tugas navigasi, timeline, review sumber, consent, aksesibilitas, dan keluar cepat. Pengamat mencatat keberhasilan, kebutuhan bantuan, waktu, kesalahan pemahaman, dan komentar.",
    )
    for text in [
        "Peserta dan tanggal: [[ISI_PESERTA_DAN_TANGGAL_UJI_PEER]].",
        "Ringkasan hasil: [[ISI_RINGKASAN_HASIL_UJI_PEER]].",
        "Masalah utama: [[ISI_MASALAH_UTAMA_UJI_PEER]].",
        "Revisi setelah pengujian: [[ISI_REVISI_SETELAH_UJI_PEER]].",
        "Masalah yang belum terselesaikan: [[ISI_SISA_MASALAH_UJI_PEER]].",
    ]:
        add_body(doc, text, first_indent=False)

    add_subheading(doc, "6.6", "Kelebihan Solusi")
    strengths = [
        "Masalah, pengguna, dan peran AI dapat dijelaskan dengan jelas dalam demonstrasi.",
        "AI dibatasi melalui sumber, schema, validator, fallback, dan human-review gate.",
        "Alur tidak langsung memaksa pengguna membuat laporan.",
        "Data sintetis dan label simulasi mengurangi risiko penggunaan data sensitif pada tugas.",
        "Prototype memiliki jalur lokal yang praktis dan tetap dapat berjalan tanpa API.",
    ]
    for text in strengths:
        add_bullet(doc, text)

    add_subheading(doc, "6.7", "Kelemahan Solusi")
    weaknesses = [
        "Belum ada validasi dengan pengguna sasaran atau penyedia layanan.",
        "Fitur keamanan produksi seperti autentikasi, enkripsi, dan audit log belum tersedia.",
        "Kualitas AI bergantung pada input, model, jaringan, dan konfigurasi.",
        "Konten edukasi dan direktori layanan belum melalui review ahli multidisipliner.",
        "Pemeriksaan aksesibilitas masih dasar dan belum mencakup seluruh teknologi asistif.",
    ]
    for text in weaknesses:
        add_bullet(doc, text)

    add_subheading(doc, "6.8", "Risiko dan Mitigasi")
    risk_rows = [
        ("Halusinasi AI", "Fakta tambahan dapat masuk draf", "Sumber wajib, null, validator, review"),
        ("Automation bias", "Pengguna menerima teks rapi tanpa mengecek", "Status pending dan keputusan eksplisit"),
        ("Kebocoran data", "Catatan sensitif terkirim/tercatat", "Minimisasi, server-side key, tanpa data nyata"),
        ("Ketergantungan layanan", "API gagal atau berubah", "Fallback dan jalur manual"),
        ("Overclaim chatbot", "Dianggap konselor/psikiater", "Label non-klinis dan handoff manusia"),
        ("Aksesibilitas semu", "Checklist teknis tidak sesuai pengalaman nyata", "Co-design dan testing pengguna"),
    ]
    add_table(doc, ["Risiko", "Dampak", "Mitigasi"], risk_rows, widths=[3.6, 5.2, 5.8], font_size=9)
    add_caption(doc, "Tabel 6.3 Risiko dan mitigasi")

    add_subheading(doc, "6.9", "Potensi Pengembangan")
    roadmap = [
        "Melakukan co-design dengan organisasi disabilitas, Satgas, pendamping, ahli trauma, dan ahli hukum.",
        "Membangun autentikasi, penyimpanan terenkripsi, consent ledger, revoke access, dan audit log.",
        "Menguji NVDA/VoiceOver, zoom, keyboard-only, switch access, caption, easy-read, dan bahasa isyarat.",
        "Memperluas dataset sintetis, red-team prompt, evaluasi bias bahasa, dan pengujian catatan panjang.",
        "Memvalidasi konten edukasi dan direktori bantuan melalui lembaga berwenang.",
        "Melakukan pilot terbatas hanya setelah penilaian etika, privasi, keamanan, dan kesiapan layanan.",
    ]
    for index, text in enumerate(roadmap, 1):
        add_bullet(doc, text, number=index)


def add_chapter_seven(doc):
    add_chapter(doc, "VII", "PENUTUP")
    add_subheading(doc, "7.1", "Kesimpulan")
    add_body(
        doc,
        "AmanAkses merupakan prototype platform digital aksesibel yang menjawab kebutuhan persiapan sebelum pengguna mencari bantuan formal. Produk menggabungkan edukasi easy-read, jurnal, bukti sintetis, kronologi bersumber, chatbot non-klinis, consent, pengaturan aksesibilitas, dan keluar cepat. Nilai utamanya adalah memberikan pilihan serta menjaga kendali pengguna.",
    )
    add_body(
        doc,
        "AI digunakan pada fungsi yang dapat dibatasi dan didemonstrasikan. Safe Timeline Assistant membantu menata catatan menjadi draf peristiwa, sedangkan Asisten Aman membantu navigasi dan persiapan. Keduanya tidak menentukan kebenaran, diagnosis, niat, kesalahan, atau keputusan hukum. Structured output, validator, fallback, allowlist, konfirmasi, dan human-review gate digunakan untuk mengurangi risiko, tetapi tidak menghapus kebutuhan tinjauan manusia.",
    )
    add_body(
        doc,
        "Validasi teknis menunjukkan build dan lint lulus, timeline melewati 10 skenario utama serta 1 negative guard, chatbot melewati 12 skenario, fallback berfungsi, dan rahasia tidak ditemukan pada bundle atau arsip. Temuan tersebut membuktikan perilaku teknis tertentu pada data sintetis, bukan kesiapan penggunaan dunia nyata. Pengujian bersama pengguna sasaran, keamanan produksi, dan validasi layanan masih harus dilakukan.",
    )

    add_subheading(doc, "7.2", "Saran")
    suggestions = [
        "Pertahankan transparansi antara fitur yang terimplementasi, simulasi UI, dan visi produksi.",
        "Jangan menggunakan data kasus nyata pada deployment demo atau pengujian kelas.",
        "Libatkan penyandang disabilitas dan penyedia layanan dalam keputusan desain berikutnya.",
        "Prioritaskan keamanan data dan tata kelola consent sebelum penyimpanan produksi.",
        "Lakukan evaluasi AI secara berkala dan jangan mengandalkan prompt sebagai satu-satunya guardrail.",
        "Isi hasil peer testing dan tautan pengumpulan hanya setelah bukti tersedia.",
    ]
    for index, text in enumerate(suggestions, 1):
        add_bullet(doc, text, number=index)


def add_references(doc):
    add_center_title(doc, "DAFTAR PUSTAKA", outline=0, page_break_before=True)
    references = [
        "Google AI for Developers. (2026). Gemini API: Structured output. https://ai.google.dev/gemini-api/docs/structured-output",
        "Kementerian Pendidikan, Kebudayaan, Riset, dan Teknologi. (2024). Peraturan Menteri Pendidikan, Kebudayaan, Riset, dan Teknologi Nomor 55 Tahun 2024 tentang Pencegahan dan Penanganan Kekerasan di Lingkungan Perguruan Tinggi. https://peraturan.bpk.go.id/Details/305767/permendikbudriset-no-55-tahun-",
        "Komisi Nasional Anti Kekerasan terhadap Perempuan. (2024). Siaran Pers Komnas Perempuan Merespons Hari Disabilitas Internasional 2024. https://komnasperempuan.go.id/siaran-pers-detail/siaran-pers-komnas-perempuan-merespons-hari-disabilitas-internasional-2024",
        "OpenAI. (2026). ChatGPT dan Codex product documentation. https://openai.com/",
        "Republik Indonesia. (2016). Undang-Undang Nomor 8 Tahun 2016 tentang Penyandang Disabilitas. https://peraturan.bpk.go.id/Home/Details/37251/uu-no-8-tahun-2016",
        "Republik Indonesia. (2022). Undang-Undang Nomor 12 Tahun 2022 tentang Tindak Pidana Kekerasan Seksual. https://peraturan.bpk.go.id/Details/207944/uu-no-12-tahun-2022",
        "Republik Indonesia. (2025). Peraturan Pemerintah Nomor 30 Tahun 2025 tentang Pencegahan Tindak Pidana Kekerasan Seksual serta Penanganan, Pelindungan, dan Pemulihan Korban. https://peraturan.bpk.go.id/Details/338353/pp-no-30-tahun-2025",
        "Tim Pengampu. (2026). Panduan Tugas Besar AI For Real Impact 2026: Solving Real Problems with AI Collaboration.",
        "World Wide Web Consortium. (2023). Web Content Accessibility Guidelines (WCAG) 2.2. https://www.w3.org/TR/WCAG22/",
        "Repository AmanAkses. (2026). Source code, fixture validasi, dokumentasi setup, diagram, dan hasil pengujian lokal.",
    ]
    for reference in references:
        paragraph = add_body(doc, reference, first_indent=False)
        paragraph.paragraph_format.left_indent = Cm(1.0)
        paragraph.paragraph_format.first_line_indent = Cm(-1.0)


def add_appendices(doc):
    add_center_title(doc, "LAMPIRAN", outline=0, page_break_before=True)

    add_subheading(doc, "Lampiran A", "Contoh Prompt dan Instruksi AI")
    add_minor_heading(doc, "A.1 Prompt Analisis Masalah")
    add_body(
        doc,
        "“Analisis ide AmanAkses berdasarkan panduan Tugas Besar AI For Real Impact. Identifikasi masalah nyata, pengguna utama, bagian yang membutuhkan AI, risiko overclaim, dan cara mendemonstrasikan solusi. Jangan menganggap semua fitur sudah tersedia.”",
        first_indent=False,
        italic=True,
    )
    add_minor_heading(doc, "A.2 Prompt Implementasi Timeline")
    add_body(
        doc,
        "“Implementasikan endpoint Gemini server-side untuk menyusun draf kronologi dari catatan sintetis terpilih. Gunakan structured JSON, sourceNoteIds wajib, field null untuk informasi hilang, requiresReview=true, dan fallback deterministik. Jangan mengirim API key ke browser.”",
        first_indent=False,
        italic=True,
    )
    add_minor_heading(doc, "A.3 Prompt Chatbot")
    add_body(
        doc,
        "“Buat Asisten Aman untuk dukungan emosional non-klinis, navigasi, dan persiapan. Batasi route internal, gunakan tool allowlist, wajibkan konfirmasi, blokir tool produktivitas pada keadaan darurat, serta larang diagnosis dan nasihat hukum.”",
        first_indent=False,
        italic=True,
    )

    add_subheading(doc, "Lampiran B", "Contoh Structured Output")
    output_example = (
        '{\n'
        '  "events": [\n'
        '    {\n'
        '      "id": "event-1",\n'
        '      "date": "2026-05-21",\n'
        '      "time": null,\n'
        '      "location": null,\n'
        '      "summary": "Pengguna berbicara dengan pendamping tepercaya.",\n'
        '      "sourceNoteIds": ["note-03"],\n'
        '      "uncertainty": "missing",\n'
        '      "requiresReview": true\n'
        '    }\n'
        '  ]\n'
        '}'
    )
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(output_example)
    set_run_font(run, size=9)

    add_subheading(doc, "Lampiran C", "Instrumen Peer Testing")
    peer_rows = [
        ("1", "Menemukan materi easy-read", "Berhasil/perlu bantuan", "Catatan"),
        ("2", "Membuat jurnal dengan data sintetis", "Berhasil/perlu bantuan", "Catatan"),
        ("3", "Menyusun timeline dari tiga catatan", "Berhasil/perlu bantuan", "Catatan"),
        ("4", "Menjelaskan sumber dan ketidakpastian", "Benar/belum benar", "Catatan"),
        ("5", "Edit, terima, dan tolak event", "Berhasil/perlu bantuan", "Catatan"),
        ("6", "Menggunakan satu tool chatbot", "Berhasil/perlu bantuan", "Catatan"),
        ("7", "Mengaktifkan aksesibilitas dan safe exit", "Berhasil/perlu bantuan", "Catatan"),
    ]
    add_table(
        doc,
        ["No.", "Tugas", "Hasil", "Observasi"],
        peer_rows,
        widths=[1.2, 6.5, 3.6, 3.3],
        font_size=9,
    )
    add_body(
        doc,
        "Skala 1–5: kemudahan, kejelasan, aksesibilitas, kepercayaan terhadap AI, kendali manusia, kegunaan timeline, dan kegunaan chatbot. Jangan meminta peserta menggunakan pengalaman pribadi atau data kasus nyata.",
        first_indent=False,
    )

    add_subheading(doc, "Lampiran D", "Tautan Artefak")
    for text in [
        "Prototype publik: [[ISI_LINK_PROTOTYPE]]",
        "Video demonstrasi: [[ISI_LINK_VIDEO_DEMO]]",
        "Presentasi PDF/PPTX: [[ISI_LINK_PRESENTASI]]",
        "Source code: [[ISI_LINK_SOURCE_CODE_ATAU_REPOSITORY]]",
        "Workbook feedback sintetis: AmanAkses_AI_For_Real_Impact/Data_Feedback_Pengguna_Sintetis.xlsx",
    ]:
        add_body(
            doc,
            text,
            first_indent=False,
            align=WD_ALIGN_PARAGRAPH.LEFT,
        )

    add_subheading(doc, "Lampiran E", "Struktur Source Code")
    structure = [
        "api/timeline.ts — endpoint Gemini untuk Safe Timeline Assistant.",
        "api/chat.ts — endpoint Gemini untuk Asisten Aman.",
        "src/lib/timelineAssistant.ts — schema, validator, dan fallback timeline.",
        "src/lib/chatAssistant.ts — schema, route allowlist, tool guard, dan fallback chatbot.",
        "src/data/mockData.ts — data sintetis.",
        "src/App.tsx — antarmuka dan human-review flow.",
        "validation/ — fixture timeline dan chatbot.",
        ".env.example — nama environment variable tanpa rahasia.",
    ]
    for text in structure:
        add_bullet(doc, text)

    add_subheading(doc, "Lampiran F", "Checklist Sebelum Pengumpulan")
    checklist = [
        "[ ] Isi seluruh placeholder [[ISI_...]] dengan bukti yang tersedia.",
        "[ ] Jalankan pnpm build, pnpm lint, pnpm validate:timeline, dan pnpm validate:chat.",
        "[ ] Pastikan .env, .env.local, node_modules, dist, dan log tidak masuk arsip.",
        "[ ] Gunakan data sintetis pada video dan demonstrasi.",
        "[ ] Perbarui daftar isi Word dengan Ctrl+A lalu F9.",
        "[ ] Buka ulang DOCX/PDF, presentasi, workbook, video, prototype, dan source archive.",
        "[ ] Jangan mengubah data sintetis menjadi klaim hasil survei nyata.",
    ]
    for text in checklist:
        add_body(doc, text, first_indent=False)


def build():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document(SOURCE)
    truncate_after_cover(doc)
    configure_styles(doc)
    add_page_numbers(doc)

    add_front_matter(doc)
    add_chapter_one(doc)
    add_chapter_two(doc)
    add_chapter_three(doc)
    add_chapter_four(doc)
    add_chapter_five(doc)
    add_chapter_six(doc)
    add_chapter_seven(doc)
    add_references(doc)
    add_appendices(doc)

    doc.core_properties.title = "Tugas Proyek Kecerdasan Buatan - Sistem AmanAkses"
    doc.core_properties.subject = "Laporan lengkap proyek AmanAkses"
    doc.core_properties.author = "Tim AmanAkses"
    doc.core_properties.keywords = "AmanAkses, AI, aksesibilitas, Gemini, human-in-the-loop"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
